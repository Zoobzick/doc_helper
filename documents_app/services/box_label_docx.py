from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document

PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")


def render_box_label_docx(template_path: str, context: dict[str, Any]) -> BytesIO:
    """
    template_path (str): путь к docx-шаблону
    context (dict): значения плейсхолдеров: {"kit_1": "...", "work_1": "...", ...}

    Возвращает BytesIO с готовым docx.
    """
    doc = Document(template_path)

    # параграфы вне таблиц
    _replace_in_paragraphs(doc.paragraphs, context)

    # таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, context)
        _cleanup_empty_rows(table)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _replace_in_paragraphs(paragraphs, context: dict[str, Any]) -> None:
    for p in paragraphs:
        for run in p.runs:
            if "{{" not in run.text:
                continue

            def repl(m: re.Match) -> str:
                key = m.group(1)
                val = context.get(key, "")
                return "" if val is None else str(val)

            # ВАЖНО: {{key}} должен быть целиком в одном run, иначе Word может не заменить
            run.text = PLACEHOLDER_RE.sub(repl, run.text)


def _cell_is_empty(cell) -> bool:
    """
    Пустая ячейка = нет текста или только пробелы.
    (У тебя в строке только {{kit_i}} -> после замены станет либо текст, либо пусто)
    """
    return not (cell.text or "").strip()


def _cleanup_empty_rows(table) -> None:
    """
    Удаляем строки таблицы, где все ячейки пустые.
    """
    to_delete = []
    for row in table.rows:
        if all(_cell_is_empty(cell) for cell in row.cells):
            to_delete.append(row)

    for row in to_delete:
        row._tr.getparent().remove(row._tr)
