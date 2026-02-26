# documents_app/services/title_sheets.py
from __future__ import annotations

import hashlib
import json
import os
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Mapping, Iterable

from django.conf import settings
from django.db import transaction

from projects_app.models import Project
from projects_app.services import sanitize_filename  # уже есть у тебя :contentReference[oaicite:1]{index=1}
from documents_app.models import TitleSheet


# (TITLE_TEMPLATE_FILENAME) имя docx-шаблона титульника в папке шаблонов
TITLE_TEMPLATE_FILENAME = "title_sheet.docx"

# (TITLE_SHEETS_DIRNAME) подпапка внутри DOCUMENTS_DIR
TITLE_SHEETS_DIRNAME = "Титульные листы"


@dataclass(frozen=True)
class TitleSheetPaths:
    # (docx_path) путь до временного/итогового docx (мы его используем для конвертации)
    docx_path: Path
    # (pdf_path) итоговый PDF, который показываем пользователю
    pdf_path: Path


def _documents_root() -> Path:
    """
    (root) корень документов, аналогично PASSPORTS_DIR/PROJECTS_DIR и т.д.
    """
    root = Path(settings.DOCUMENTS_DIR).resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def _title_sheets_root() -> Path:
    """
    (dir_) папка для титульных листов: DOCUMENTS_DIR / "Титульные листы"
    """
    dir_ = (_documents_root() / TITLE_SHEETS_DIRNAME).resolve()
    dir_.mkdir(parents=True, exist_ok=True)
    return dir_


def _template_path() -> Path:
    """
    (path) полный путь к title_sheet.docx
    Ожидаем, что шаблоны лежат в settings.DOCX_TEMPLATES_DIR (как у тебя в documents_app для box_label). :contentReference[oaicite:2]{index=2}
    """
    base = Path(settings.DOCX_TEMPLATES_DIR).resolve()
    path = (base / TITLE_TEMPLATE_FILENAME).resolve()
    if not path.exists():
        raise RuntimeError(f"Шаблон титульника не найден: {path}")
    return path


def _sha256_bytes(data: bytes) -> str:
    # (h) sha256-хеш содержимого
    h = hashlib.sha256()
    h.update(data)
    return h.hexdigest()


def _sha256_file(path: Path) -> str:
    # (h) sha256-хеш файла на диске
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _json_dumps_stable(obj: Any) -> bytes:
    """
    (b) стабильная сериализация для хеширования:
    - сортируем ключи
    - без лишних пробелов
    - UTF-8
    """
    s = json.dumps(obj, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return s.encode("utf-8")


def _doc_label_for_type(doc_type: str) -> str:
    """
    (label) строка, которая пойдёт в {{documentation}}
    """
    if doc_type == TitleSheet.DocType.ID:
        return "Исполнительная документация"
    if doc_type == TitleSheet.DocType.RD:
        return "Рабочая документация"
    if doc_type == TitleSheet.DocType.ID_RD:
        return "Исполнительная и рабочая документация"
    raise ValueError(f"Неизвестный doc_type: {doc_type}")


def build_title_placeholders(*, project: Project, doc_type: str) -> dict[str, str]:
    """
    Возвращает словарь плейсхолдеров для замены в docx.
    Все источники — из справочников проекта (без парсинга).
    """
    # (line_full) полное наименование линии
    line_full = project.line.full_name if project.line else ""
    # (stage_full) полное наименование этапа
    stage_full = project.stage.full_name if project.stage else ""
    # (plot_full) полное наименование участка
    plot_full = project.plot.full_name if project.plot else ""
    # (construction_full) полное наименование конструкции (у тебя это TextField)
    construction_full = project.construction or ""
    # (documentation) строка для {{documentation}}
    documentation = _doc_label_for_type(doc_type)
    # (full_code) шифр проекта
    full_code = project.full_code or ""

    return {
        "{{project_line_full}}": line_full,
        "{{project_stage_full}}": stage_full,
        "{{project_plot_full}}": plot_full,
        "{{project_construction_full}}": construction_full,
        "{{documentation}}": documentation,
        "{{project_full_code}}": full_code,
    }


def compute_context_sha256(placeholders: Mapping[str, Any]) -> str:
    """
    Хешируем именно "данные для подстановки", чтобы:
    - если меняется хоть одно поле в справочниках/проекте → пересборка
    """
    # (payload) берём только значения; ключи тоже важны, но они стабильны
    payload = dict(placeholders)
    return _sha256_bytes(_json_dumps_stable(payload))


def compute_template_sha256(template_path: Path) -> str:
    """
    Хешируем байты шаблона docx.
    Любое изменение шаблона → пересборка.
    """
    return _sha256_file(template_path)


def _build_paths(*, project: Project, doc_type: str) -> TitleSheetPaths:
    """
    Строим имена файлов по твоему требованию:
    DOCUMENTS_DIR / "Титульные листы" / "Титул {documentation} {project_full_code}.pdf"
    """
    root = _title_sheets_root()

    # (documentation) человекочитаемая строка
    documentation = _doc_label_for_type(doc_type)
    # (base_name) базовое имя без расширения
    base_name = f"Титул {documentation} {project.full_code or ''}".strip()

    safe_stem = sanitize_filename(base_name)
    pdf_path = (root / f"{safe_stem}.pdf").resolve()
    docx_path = (root / f"{safe_stem}.docx").resolve()

    return TitleSheetPaths(docx_path=docx_path, pdf_path=pdf_path)


def _find_soffice_cmd() -> str:
    """
    На Linux обычно доступно 'libreoffice' или 'soffice'.
    """
    for candidate in ("libreoffice", "soffice"):
        if shutil_which(candidate):
            return candidate
    # fallback: пусть subprocess попробует (вдруг в PATH нестандартно)
    return "libreoffice"


def shutil_which(name: str) -> str | None:
    # локальная мини-версия shutil.which, чтобы не тащить импорт везде
    for p in os.environ.get("PATH", "").split(os.pathsep):
        cand = Path(p) / name
        if cand.exists() and os.access(str(cand), os.X_OK):
            return str(cand)
    return None


def convert_docx_to_pdf(*, docx_path: Path, out_dir: Path) -> Path:
    """
    Конвертация DOCX → PDF через LibreOffice headless.
    Возвращаем путь к получившемуся PDF (в out_dir).
    """
    out_dir.mkdir(parents=True, exist_ok=True)

    cmd = _find_soffice_cmd()
    # (args) параметры LibreOffice
    args = [
        cmd,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(docx_path),
    ]

    proc = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if proc.returncode != 0:
        raise RuntimeError(
            "LibreOffice не смог сконвертировать DOCX в PDF.\n"
            f"cmd: {' '.join(args)}\n"
            f"stdout: {proc.stdout}\n"
            f"stderr: {proc.stderr}"
        )

    produced = (out_dir / f"{docx_path.stem}.pdf").resolve()
    if not produced.exists():
        raise RuntimeError(f"PDF после конвертации не найден: {produced}")

    return produced


# ====== ТВОЙ РАБОЧИЙ КОД ЗАМЕНЫ ПЛЕЙСХОЛДЕРОВ (почти без изменений) ======

from docx import Document  # noqa: E402


def _replace_placeholders_in_paragraph(paragraph, placeholders: Mapping[str, Any]) -> None:
    """
    paragraph  – объект Paragraph.
    placeholders – {"{{name}}": "value", ...}

    Логика:
      1. Склеиваем текст всех runs в одну строку (full_text).
      2. Находим плейсхолдеры в full_text.
      3. Для каждого вхождения:
         - по карте index_map вычисляем, в каких run-ах лежит;
         - берём формат первого run плейсхолдера;
         - в этот run записываем: before + value + after;
         - промежуточные run-ы очищаем (run.text = "").
    """
    if not paragraph.runs:
        return

    def build_index():
        full = ""
        index = []
        for run_idx, run in enumerate(paragraph.runs):
            text = run.text or ""
            full += text
            index.extend((run_idx, i) for i in range(len(text)))
        return full, index

    full_text, index_map = build_index()
    if not full_text:
        return

    for placeholder, raw_value in placeholders.items():
        placeholder = str(placeholder)
        value = str(raw_value)

        while True:
            pos = full_text.find(placeholder)
            if pos == -1:
                break

            start = pos
            end = pos + len(placeholder)

            if end - 1 >= len(index_map):
                break

            run_start_idx, off_start = index_map[start]
            run_end_idx, off_end = index_map[end - 1]

            run_start = paragraph.runs[run_start_idx]
            run_end = paragraph.runs[run_end_idx]

            before = (run_start.text or "")[:off_start]
            after = (run_end.text or "")[off_end + 1:]

            run_start.text = before + value + after

            for i in range(run_start_idx + 1, run_end_idx + 1):
                paragraph.runs[i].text = ""

            full_text, index_map = build_index()


def _process_table(table, placeholders: Mapping[str, Any]) -> None:
    """
    Рекурсивно обходит таблицу и все вложенные таблицы в её ячейках.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, placeholders)

            for nested_table in cell.tables:
                _process_table(nested_table, placeholders)


def render_title_docx(*, template_path: Path, output_path: Path, placeholders: Mapping[str, Any]) -> None:
    """
    template_path – исходный papka_title.docx
    output_path   – куда сохранить результат docx
    """
    doc = Document(str(template_path))

    for paragraph in doc.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, placeholders)

    for table in doc.tables:
        _process_table(table, placeholders)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ====== ГЛАВНЫЕ ФУНКЦИИ ДЛЯ ВЬЮХ/PROJECT_DETAIL ======

@transaction.atomic
def ensure_title_sheet(*, project: Project, doc_type: str) -> TitleSheet:
    """
    Гарантирует, что для (project, doc_type) есть актуальный PDF (если не locked).
    Возвращает запись TitleSheet.
    """
    # (template_path) путь до шаблона
    template_path = _template_path()
    # (placeholders) данные для подстановки
    placeholders = build_title_placeholders(project=project, doc_type=doc_type)

    # (context_sha) хеш данных
    context_sha = compute_context_sha256(placeholders)
    # (template_sha) хеш шаблона
    template_sha = compute_template_sha256(template_path)

    obj, _ = TitleSheet.objects.select_for_update().get_or_create(
        project=project,
        doc_type=doc_type,
        defaults={
            "context_sha256": "",
            "template_sha256": "",
            "pdf_path": "",
            "docx_path": "",
            "is_locked": False,
        },
    )

    paths = _build_paths(project=project, doc_type=doc_type)

    # (pdf_exists) существует ли текущий PDF на диске
    pdf_exists = bool(obj.pdf_path) and Path(obj.pdf_path).exists()

    # ✅ Если заморожено и PDF есть — ничего не трогаем, даже если шаблон/данные изменились.
    if obj.is_locked and pdf_exists:
        return obj

    needs_regen = (
        (not pdf_exists)
        or (obj.context_sha256 != context_sha)
        or (obj.template_sha256 != template_sha)
        or (not obj.pdf_path)
    )

    if not needs_regen:
        return obj

    # 1) генерим docx (в ту же папку, что и pdf)
    render_title_docx(
        template_path=template_path,
        output_path=paths.docx_path,
        placeholders=placeholders,
    )

    # 2) конвертим docx → pdf
    produced_pdf = convert_docx_to_pdf(docx_path=paths.docx_path, out_dir=paths.pdf_path.parent)

    # 3) переименовываем в ожидаемое имя (на всякий случай)
    if produced_pdf != paths.pdf_path:
        produced_pdf.replace(paths.pdf_path)

    # 4) сохраняем в БД пути и хеши
    obj.pdf_path = str(paths.pdf_path)
    obj.docx_path = str(paths.docx_path)  # сейчас не обязателен, но пусть будет
    obj.context_sha256 = context_sha
    obj.template_sha256 = template_sha
    obj.save(update_fields=["pdf_path", "docx_path", "context_sha256", "template_sha256", "updated_at"])

    return obj


def ensure_all_title_sheets_for_project(project: Project) -> None:
    """
    Обновляет все 3 титульника (если не locked) — удобно дергать при заходе на project_detail.
    """
    for dt in (TitleSheet.DocType.ID, TitleSheet.DocType.RD, TitleSheet.DocType.ID_RD):
        ensure_title_sheet(project=project, doc_type=dt)