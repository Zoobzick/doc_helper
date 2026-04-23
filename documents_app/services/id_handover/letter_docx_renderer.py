from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt

from documents_app.models import DocumentBatchLetterType


FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(13)


class LetterDocxRendererError(Exception):
    """Базовая ошибка рендера письма."""


class LetterDocxRendererValidationError(LetterDocxRendererError):
    """Ошибка структуры DOCX-шаблона письма."""


class LetterDocxRenderer:
    """
    Рендерит DOCX-письмо по готовому context.

    Логика построена по рабочему шаблонному подходу:
    - шапка заменяется глобально
    - список проектов строится копированием строки-образца
    - приложение строится копированием блока строк:
      - 2 строки на project для FOR_EXECUTION
      - 3 строки на project для TO_ARCHIVE
    """

    COMPOSITION_SENTINEL = "{{project}}. {{plot}}. {{construction}}"
    APPENDIX_SENTINEL = "Реестр исполнительной документации по шифру {{project}}"

    def render_from_context(
        self,
        *,
        context: dict[str, Any],
        template_path: str | Path,
        output_path: str | Path,
    ) -> Path:
        template_path = Path(template_path)
        output_path = Path(output_path)

        if not template_path.exists():
            raise LetterDocxRendererValidationError(
                f"Шаблон письма не найден: {template_path}"
            )

        document = Document(str(template_path))

        global_replacements = self._build_global_replacements(context=context)

        # 1. Шапка письма
        for paragraph in document.paragraphs:
            self._replace_tokens_in_paragraph(paragraph, global_replacements)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_tokens_in_paragraph(paragraph, global_replacements)

        # 2. Список проектов
        self._render_projects_list(
            document=document,
            context=context,
        )

        # 3. Приложение
        self._render_appendix(
            document=document,
            context=context,
            global_replacements=global_replacements,
        )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path

    def _build_global_replacements(self, *, context: dict[str, Any]) -> dict[str, str]:
        placeholders = context.get("placeholders", {})

        # Поддерживаем и «dd», и dd — на случай расхождения шаблонов.
        return {
            "{{number}}": str(placeholders.get("number", "") or ""),
            "«dd»": str(placeholders.get("dd", "") or ""),
            "dd": str(placeholders.get("dd", "") or ""),
            "mm": str(placeholders.get("mm", "") or ""),
            "yyyy": str(placeholders.get("yyyy", "") or ""),
            "{{period}}": str(placeholders.get("period", "") or ""),
            "{{project_line_full_name}}": str(placeholders.get("project_line_full_name", "") or ""),
            "{{project_line_code}}": str(placeholders.get("project_line_code", "") or ""),
        }

    def _render_projects_list(
        self,
        *,
        document: Document,
        context: dict[str, Any],
    ) -> None:
        projects = context.get("projects", [])
        if not projects:
            raise LetterDocxRendererValidationError(
                "Контекст письма не содержит projects."
            )

        idx = self._find_paragraph_index_by_contains(
            document=document,
            needle=self.COMPOSITION_SENTINEL,
        )
        if idx is None:
            raise LetterDocxRendererValidationError(
                f"Не найден абзац со строкой: {self.COMPOSITION_SENTINEL}"
            )

        base_paragraph = document.paragraphs[idx]

        def format_project_line(item: dict[str, Any]) -> str:
            project = str(item.get("project", "") or "")
            plot = str(item.get("plot", "") or "")
            construction = str(item.get("construction", "") or "")
            return f"{project}. {plot}. {construction}".strip()

        self._set_paragraph_text(base_paragraph, format_project_line(projects[0]))

        last_paragraph = base_paragraph
        for item in projects[1:]:
            new_paragraph = self._insert_paragraph_after(last_paragraph)
            self._set_paragraph_text(new_paragraph, format_project_line(item))
            last_paragraph = new_paragraph

    def _render_appendix(
        self,
        *,
        document: Document,
        context: dict[str, Any],
        global_replacements: dict[str, str],
    ) -> None:
        appendix_items = context.get("appendix", {}).get("items", [])
        if not appendix_items:
            raise LetterDocxRendererValidationError(
                "Контекст письма не содержит appendix.items."
            )

        batch = context.get("batch", {})
        letter_type = batch.get("letter_type")
        lines_per_project = self._resolve_lines_per_project(letter_type=letter_type)

        if len(appendix_items) % lines_per_project != 0:
            raise LetterDocxRendererValidationError(
                "Количество appendix.items не кратно числу строк приложения на проект."
            )

        idx_app1 = self._find_paragraph_index_by_contains(
            document=document,
            needle=self.APPENDIX_SENTINEL,
        )
        if idx_app1 is None:
            raise LetterDocxRendererValidationError(
                f"Не найден блок приложения: абзац '{self.APPENDIX_SENTINEL}'"
            )

        paragraphs = document.paragraphs
        if idx_app1 + lines_per_project - 1 >= len(paragraphs):
            raise LetterDocxRendererValidationError(
                "Шаблон обрывается: не хватает строк шаблонного блока приложения."
            )

        template_block = [
            paragraphs[idx_app1 + offset]
            for offset in range(lines_per_project)
        ]

        template_texts = [
            self._paragraph_text(paragraph)
            for paragraph in template_block
        ]

        projects_groups = [
            appendix_items[i:i + lines_per_project]
            for i in range(0, len(appendix_items), lines_per_project)
        ]

        # Первый проект — в существующий шаблонный блок
        for row_index, item in enumerate(projects_groups[0]):
            self._apply_appendix_item_to_paragraph(
                paragraph=template_block[row_index],
                template_text=template_texts[row_index],
                item=item,
                global_replacements=global_replacements,
            )

        last_block_last_paragraph = template_block[-1]

        # Остальные проекты — вставляем копии блока
        for project_items in projects_groups[1:]:
            new_block: list[Paragraph] = []

            for row_index, item in enumerate(project_items):
                new_paragraph = self._insert_paragraph_after(last_block_last_paragraph)
                self._set_paragraph_text(new_paragraph, template_texts[row_index])
                self._apply_appendix_item_to_paragraph(
                    paragraph=new_paragraph,
                    template_text=template_texts[row_index],
                    item=item,
                    global_replacements=global_replacements,
                )
                new_block.append(new_paragraph)
                last_block_last_paragraph = new_paragraph

    def _apply_appendix_item_to_paragraph(
        self,
        *,
        paragraph: Paragraph,
        template_text: str,
        item: dict[str, Any],
        global_replacements: dict[str, str],
    ) -> None:
        text = template_text

        replacements = {
            **global_replacements,
            "{{project}}": str(item.get("project", "") or ""),
            "{{plot}}": str(item.get("plot", "") or ""),
            "{{construction}}": str(item.get("construction", "") or ""),
            "{{pages}}": str(item.get("pages_count", "") or ""),
            "{{pages_str}}": str(item.get("pages_str", "") or ""),
        }

        # Если в контексте уже передан готовый текст строки — используем его напрямую.
        ready_text = str(item.get("text", "") or "").strip()

        if ready_text:
            # Нумерация приложений уже задана в DOCX-шаблоне списком Word.
            # Не дублируем её вручную в тексте строки.
            self._set_paragraph_text(paragraph, ready_text)
            return

        for token, value in replacements.items():
            text = text.replace(token, value)

        self._set_paragraph_text(paragraph, text)

    def _resolve_lines_per_project(self, *, letter_type: str | None) -> int:
        if letter_type == DocumentBatchLetterType.FOR_EXECUTION:
            return 2
        if letter_type == DocumentBatchLetterType.TO_ARCHIVE:
            return 3
        raise LetterDocxRendererValidationError(
            f"Неподдерживаемый тип письма: {letter_type}"
        )

    def _paragraph_text(self, paragraph: Paragraph) -> str:
        return "".join(run.text for run in paragraph.runs)

    def _set_paragraph_text(self, paragraph: Paragraph, text: str) -> None:
        paragraph.clear()
        run = paragraph.add_run(text)
        self._apply_font(run)

    def _apply_font(self, run) -> None:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE

        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), FONT_NAME)
        r_fonts.set(qn("w:hAnsi"), FONT_NAME)
        r_fonts.set(qn("w:cs"), FONT_NAME)
        r_fonts.set(qn("w:eastAsia"), FONT_NAME)

    def _replace_tokens_in_paragraph(
        self,
        paragraph: Paragraph,
        replacements: dict[str, str],
    ) -> None:
        full_text = self._paragraph_text(paragraph)
        if not full_text:
            return

        new_text = full_text
        for token, value in replacements.items():
            new_text = new_text.replace(token, value)

        if new_text != full_text:
            self._set_paragraph_text(paragraph, new_text)

    def _find_paragraph_index_by_contains(
        self,
        *,
        document: Document,
        needle: str,
    ) -> int | None:
        for index, paragraph in enumerate(document.paragraphs):
            if needle in self._paragraph_text(paragraph):
                return index
        return None

    def _insert_paragraph_after(self, paragraph: Paragraph) -> Paragraph:
        new_p_xml = deepcopy(paragraph._p)

        for text_node in new_p_xml.xpath(".//w:t"):
            text_node.text = ""

        paragraph._p.addnext(new_p_xml)
        return Paragraph(new_p_xml, paragraph._parent)
