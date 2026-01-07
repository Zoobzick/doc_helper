import os

from docx import Document
from openpyxl.workbook import Workbook
from django.conf import settings

class CheckService:

    def __init__(self,
                 month: int | str,
                 year: int = 2025
                 ):
        self.year = year
        self.month = month
        self.base_path = settings.BASE_ID_DIR
        self.months = {'январь': '01',
                       'февраль': '02',
                       'март': '03',
                       'апрель': '04',
                       'май': '05',
                       'июнь': '06',
                       'июль': '07',
                       'август': '08',
                       'сентябрь': '09',
                       'октябрь': '10',
                       'ноябрь': '11',
                       'декабрь': '12'}
        self.full_path = self._build_full_path()

    def _build_full_path(self):
        """Корректно строит полный путь к папке месяца"""
        month_num = ""
        month_name = ""

        if self.month.isdigit():
            # Для цифрового ввода
            month_num = f"{int(self.month):02d}"
            # Ищем название месяца
            for name, num in self.months.items():
                if num == self.month:
                    month_name = name.capitalize()
                    break
        else:
            # Для текстового ввода
            month_lower = self.month.lower()
            if month_lower in self.months:
                month_num = f"{int(self.months[month_lower]):02d}"
                month_name = self.month.capitalize()
            else:
                raise ValueError(f"Неизвестный месяц: {self.month}")
        year_path = os.path.join(self.base_path, str(self.year))
        self.month = f"{month_num}. {month_name}"
        # Используем os.path.join для корректного построения пути
        return os.path.join(year_path, f"{month_num}. {month_name}")

    @staticmethod
    def clean_project_name(name):
        """Очищает название проекта от лишних символов"""
        # Удаляем знаки препинания в начале и конце
        import string
        return name.strip(string.punctuation + ' ')

    def get_list_of_files(self, path=None):
        """
        Поиск файлов в корневом каталоге и субкаталоге
        """
        if path is None:
            path = self.full_path  # Первый вызов - используем full_path

        current_dir_files = os.listdir(path)  # ← Используем переданный путь
        all_files = list()

        for el in current_dir_files:
            full_path = os.path.join(path, el)  # ← Используем переданный путь

            if os.path.isdir(full_path):
                # Рекурсивно вызываем для ПОДПАПКИ
                all_files = all_files + self.get_list_of_files(full_path)  # ← ПЕРЕДАЕМ ПУТЬ!
            else:
                all_files.append(full_path)

        return CheckService.docx_separator(all_files)

    @staticmethod
    def docx_separator(file_list) -> list:
        """
        Выделяет из списка всех файлов только файлы с расширением ".docx"
        :param file_list: список файлов в корневом каталоге и субкаталоге
        :return:
        """
        result = []
        for el in file_list:
            filename = os.path.basename(el)

            is_docx = el.endswith("docx")
            start_with_act = filename.startswith("Акт")

            if is_docx and start_with_act:
                result.append(el)
            else:
                pass
        return result

    @staticmethod
    def get_projects_and_path(docx_files, test_mark="ИМИП-МРАЛ"):
        test_mark = test_mark
        result = {}

        for pth in docx_files:
            doc = Document(pth)
            normalized_path = str(pth).lower()
            if test_mark in doc.tables[2].rows[0].cells[0].text:
                for word in doc.tables[2].rows[0].cells[0].text.split():
                    clean_word = CheckService.clean_project_name(word)
                    if test_mark in clean_word:
                        result[normalized_path] = clean_word.lower()
                    else:
                        pass
            else:
                pass
        return result

    @staticmethod
    def check_project_and_path(path_and_project: dict, verbose=True):
        """
        Проверяет, содержится ли название проекта в пути к файлу.

        Args:
            path_and_project: Словарь {путь_к_файлу: название_проекта}
            verbose: Если True, выводит подробный отчет

        Returns:
            Словарь с результатами проверки
        """
        results = {
            'total': 0,
            'matches': 0,
            'mismatches': 0,
            'details': []
        }

        for file_path, project_name in path_and_project.items():
            results['total'] += 1

            # Проверяем наличие проекта в пути
            if project_name and project_name in file_path:
                results['matches'] += 1
                status = "СООТВЕТСТВУЕТ"
                is_match = True
            else:
                results['mismatches'] += 1
                status = "НЕ СООТВЕТСТВУЕТ"
                is_match = False

            # Формируем детальную информацию
            detail = {
                'file_path': file_path,
                'project_name': project_name,
                'status': status,
                'is_match': is_match
            }
            results['details'].append(detail)

        # Вывод результатов
        if verbose:
            CheckService._print_check_results(results)

        return results

    @staticmethod
    def _print_check_results(results):
        """
        Выводит результаты проверки в консоль.
        """
        print("\n" + "=" * 80)
        print("РЕЗУЛЬТАТЫ ПРОВЕРКИ СООТВЕТСТВИЯ ПУТЕЙ И ПРОЕКТОВ")
        print("=" * 80)
        print(f"Всего проверено файлов: {results['total']}")
        print(f"Соответствуют: {results['matches']}")
        print(f"Не соответствуют: {results['mismatches']}")

        if results['mismatches'] > 0:
            print("\n" + "-" * 80)
            print("ДЕТАЛИ ПО ФАЙЛАМ С НЕСООТВЕТСТВИЯМИ:")
            print("-" * 80)

            for detail in results['details']:
                if not detail['is_match']:
                    print(f"\nФайл: {detail['file_path']}")
                    print(f"Проект из акта: {detail['project_name']}")
                    print(f"Статус: {detail['status']}")

                    # Предлагаем рекомендацию
                    if detail['project_name']:
                        print(f"Рекомендация: Переместите файл в папку, содержащую '{detail['project_name']}'")

        print("\n" + "=" * 80)
        print("ПРОВЕРКА ЗАВЕРШЕНА")
        print("=" * 80)

    @staticmethod
    def find_act_doc_files(start_directory):
        """
        Находит все файлы начинающиеся на 'Акт' с расширением .doc в каталоге и подкаталогах

        Args:
            start_directory (str): Путь к начальной директории для поиска

        Returns:
            list: Список полных путей к найденным файлам
        """
        act_doc_files = []

        print(f"🔍 Поиск файлов 'Акт*.doc' в директории: {start_directory}")

        # Рекурсивно обходим все подкаталоги
        for root, dirs, files in os.walk(start_directory):
            for file in files:
                # Проверяем, что файл начинается на 'Акт' и имеет расширение .doc
                if file.startswith('Акт') and file.lower().endswith('.doc'):
                    full_path = os.path.join(root, file)
                    act_doc_files.append(full_path)
                    print(f"📄 Найден: {file}")

        print(f"✅ Найдено файлов: {len(act_doc_files)}")
        return act_doc_files

    @staticmethod
    def create_excel_from_data(data):
        """
        Создает Excel файл с двумя столбцами из данных
        и сохраняет по указанному пути

        Args:
            data: словарь с данными {'МЩК': set(), 'Тупики': set()}
        """
        # Создаем новую рабочую книгу
        wb = Workbook()
        ws = wb.active
        ws.title = "Данные"

        # Задаем заголовки столбцов
        ws['A1'] = 'МЩК'
        ws['B1'] = 'Тупики'

        # Получаем данные в виде списков и сортируем их
        mshk_list = sorted(list(data['МЩК']))
        tupiki_list = sorted(list(data['Тупики']))

        # Определяем максимальную длину для итерации
        max_len = max(len(mshk_list), len(tupiki_list))

        # Заполняем данные
        for i in range(max_len):
            row_num = i + 2  # +2 потому что первая строка - заголовки

            # Заполняем столбец МЩК
            if i < len(mshk_list):
                ws.cell(row=row_num, column=1, value=mshk_list[i])

            # Заполняем столбец Тупики
            if i < len(tupiki_list):
                ws.cell(row=row_num, column=2, value=tupiki_list[i])

        # Настраиваем ширину столбцов
        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 50

        # Указанный путь для сохранения
        save_path = r"\\Marina\ид участок №5 (липовая роща)\2025\10. Октябрь\Реестры\данные_мщк_тупики.xlsx"

        # Сохраняем файл
        wb.save(save_path)
        print(f"Файл сохранен: {save_path}")
        print(f"Количество записей: МЩК - {len(mshk_list)}, Тупики - {len(tupiki_list)}")

        return save_path

    @staticmethod
    def projects_list_excel_file(data):

        # Создаем Excel файл
        CheckService.create_excel_from_data(data)
