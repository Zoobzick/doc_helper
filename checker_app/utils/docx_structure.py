# from dataclasses import dataclass
# from pathlib import Path
# from typing import List, Tuple
#
# from docx import Document
#
#
# @dataclass
# class ParagraphItem:
#     index: int          # (index) — номер параграфа по порядку (0..)
#     text: str           # (text) — очищенный текст параграфа
#
#
# @dataclass
# class TableCellItem:
#     table_index: int    # (table_index) — номер таблицы по порядку (0..)
#     row: int            # (row) — индекс строки
#     col: int            # (col) — индекс колонки
#     text: str           # (text) — текст ячейки
#
#
# def _clean(text: str) -> str:
#     # (text) — исходный текст
#     return " ".join(text.replace("\xa0", " ").split())
#
#
# def extract_docx_structure(docx_path: Path) -> Tuple[List[ParagraphItem], List[TableCellItem]]:
#     # (docx_path) — путь до .docx
#     doc = Document(str(docx_path))  # (doc) — документ python-docx
#
#     paragraphs: List[ParagraphItem] = []
#     for i, p in enumerate(doc.paragraphs):  # (i) — номер параграфа, (p) — объект Paragraph
#         txt = _clean(p.text)                # (txt) — текст параграфа
#         if txt:                             # пустые строки пропускаем
#             paragraphs.append(ParagraphItem(index=i, text=txt))
#
#     cells: List[TableCellItem] = []
#     for ti, table in enumerate(doc.tables):                 # (ti) — номер таблицы
#         for ri, row in enumerate(table.rows):              # (ri) — номер строки
#             for ci, cell in enumerate(row.cells):          # (ci) — номер колонки
#                 txt = _clean(cell.text)                    # (txt) — текст ячейки
#                 if txt:
#                     cells.append(TableCellItem(ti, ri, ci, txt))
#
#     return paragraphs, cells

from dataclasses import dataclass
from typing import List, Tuple, Union, IO

from docx import Document


@dataclass
class ParagraphItem:
    index: int
    text: str


@dataclass
class TableCellItem:
    table_index: int
    row: int
    col: int
    text: str


def _clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def extract_docx_structure(source: Union[str, "bytes", IO]) -> Tuple[List[ParagraphItem], List[TableCellItem]]:
    """
    (source) — либо путь до файла, либо file-like объект (UploadedFile.file).
    """
    doc = Document(source)

    paragraphs: List[ParagraphItem] = []
    for i, p in enumerate(doc.paragraphs):
        txt = _clean(p.text)
        if txt:
            paragraphs.append(ParagraphItem(index=i, text=txt))

    cells: List[TableCellItem] = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = _clean(cell.text)
                if txt:
                    cells.append(TableCellItem(ti, ri, ci, txt))

    return paragraphs, cells
