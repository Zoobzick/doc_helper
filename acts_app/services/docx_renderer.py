# acts_app/services/docx_renderer.py
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Any, Optional

from docx import Document  # python-docx
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from acts_app.services.act_context import ActContext


@dataclass(frozen=True)
class RenderResult:
    """
    bytes — готовый DOCX
    filename — рекомендуемое имя файла
    """
    content: bytes
    filename: str


class ActDocxRenderer:
    """
    Первый рабочий рендер DOCX.
    Цель шага: “верно по смыслу и порядку”, а не “идеально по стилю”.

    Потом можно будет:
    - перейти на DOCX-шаблон (placeholder-ы)
    - добавить фирменные стили, таблицы, отступы, шрифты
    """

    def render(self, ctx: ActContext) -> RenderResult:
        data = ctx.data  # data (dict)

        act = data["act"]  # act (dict)
        project_codes = data["project_codes"]  # project_codes (list[str])
        materials = data["materials"]  # materials (list[dict])
        quality = data["quality_docs"]  # quality (dict)
        appendix = data["appendix"]  # appendix (dict)

        doc = Document()

        # --- базовый стиль (очень мягко) ---
        self._set_default_font(doc, name="Times New Roman", size_pt=12)

        # Заголовок
        doc.add_paragraph("АКТ ОСВИДЕТЕЛЬСТВОВАНИЯ СКРЫТЫХ РАБОТ").runs[0].bold = True
        doc.add_paragraph(f"№ {act['number']} от {act['act_date_str']}")

        # Шифры проектов
        if project_codes:
            doc.add_paragraph(f"Шифр(ы) проектной документации: {', '.join(project_codes)}")

        doc.add_paragraph("")  # пустая строка

        # 1. Наименование работ
        doc.add_paragraph("1. Наименование работ:").runs[0].bold = True
        doc.add_paragraph(act["work_name"] or "")

        # 2. Материалы / паспорта
        doc.add_paragraph("")
        doc.add_paragraph("2. При выполнении работ применены материалы и паспорта на них:").runs[0].bold = True
        if not materials:
            doc.add_paragraph("—")
        else:
            for m in materials:
                # m (dict)
                mat_name = (m.get("material_name") or "").strip()
                doc_name = (m.get("document_name") or "").strip()
                doc_date = (m.get("document_date_str") or "").strip()
                sheets = m.get("sheets_count", 1)

                line = f"{mat_name} — {doc_name}"
                if doc_date:
                    line += f" от {doc_date}"
                line += f" ({sheets} л.)"

                p = doc.add_paragraph(style=None)
                p.add_run("• ")
                p.add_run(line)

        # 3. Документы соответствия (схема всегда первая)
        doc.add_paragraph("")
        doc.add_paragraph(
            "3. Предъявлены документы, подтверждающие соответствие работ предъявляемым к ним требованиям:"
        ).runs[0].bold = True

        # 3.1 Исполнительная схема (всегда первая)
        scheme = quality["exec_scheme"]  # scheme (dict)
        scheme_line = f"{scheme['title']} {scheme['scheme_name']}"
        if scheme.get("doc_date_str"):
            scheme_line += f" от {scheme['doc_date_str']}"
        scheme_line += f" ({scheme.get('sheets_count', 1)} л.)"
        p = doc.add_paragraph()
        p.add_run("• ")
        p.add_run(scheme_line)

        # 3.2 Остальные документы
        docs = quality["docs"]  # docs (list[dict])
        if not docs:
            doc.add_paragraph("• (прочих документов нет)")
        else:
            for d in docs:
                title = (d.get("title") or "").strip()
                doc_no = (d.get("doc_no") or "").strip()
                doc_date = (d.get("doc_date_str") or "").strip()
                sheets = d.get("sheets_count", 1)

                line = title
                if doc_no:
                    line += f" №{doc_no}"
                if doc_date:
                    line += f" от {doc_date}"
                line += f" ({sheets} л.)"

                p = doc.add_paragraph()
                p.add_run("• ")
                p.add_run(line)

        # 4. Даты
        doc.add_paragraph("")
        doc.add_paragraph("4. Сроки выполнения работ:").runs[0].bold = True
        doc.add_paragraph(f"Дата начала работ: {act['work_start_date_str'] or '—'}")
        doc.add_paragraph(f"Дата окончания работ: {act['work_end_date_str'] or '—'}")

        # 5. Нормативка (пока текстом)
        doc.add_paragraph("")
        doc.add_paragraph("5. Работы выполнены в соответствии с:").runs[0].bold = True
        doc.add_paragraph(act["work_norms_text"] or "—")

        # 6. Последующие работы
        doc.add_paragraph("")
        doc.add_paragraph("6. Разрешается производство последующих работ:").runs[0].bold = True
        doc.add_paragraph(act["allow_next_works_text"] or "—")

        # 7. Доп сведения
        doc.add_paragraph("")
        doc.add_paragraph("7. Дополнительные сведения:").runs[0].bold = True
        doc.add_paragraph(act["extra_info_text"] or "—")

        # 8. Экземпляры
        doc.add_paragraph("")
        doc.add_paragraph(f"Акт составлен в {act['copies_count']} экземплярах.")

        # Приложения (если уже собрались AppendixBuilder-ом)
        doc.add_paragraph("")
        doc.add_paragraph("Приложения:").runs[0].bold = True

        lines = appendix.get("lines") or []
        if not lines:
            doc.add_paragraph("—")
        else:
            for ln in lines:
                label = ln["label"]
                sheets = ln["sheets_count"]
                p = doc.add_paragraph()
                p.add_run("• ")
                p.add_run(f"{label} ({sheets} л.)")

        if appendix.get("error"):
            # заметка для разработчика/оператора (можно убрать позже)
            doc.add_paragraph("")
            warn = doc.add_paragraph("Внимание: приложения собраны не полностью/с ошибкой.")
            warn.runs[0].bold = True
            doc.add_paragraph(appendix["error"])

        # --- output bytes ---
        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        safe_number = str(act["number"]).replace("/", "-").replace("\\", "-")
        filename = f"Акт_{safe_number}.docx"

        return RenderResult(content=content, filename=filename)

    # ---------------- helpers ----------------

    def _set_default_font(self, doc: Document, *, name: str, size_pt: int) -> None:
        style = doc.styles["Normal"]
        font = style.font
        font.name = name
        font.size = Pt(size_pt)
