# acts_app/services/act_docx_generator.py
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional, List

from django.conf import settings
from docx import Document
from docx.shared import Pt

from acts_app.models import Act
from acts_app.services.appendix_builder import AppendixBuilder
from acts_app.services.act_docx_context import build_act_docx_context


class DocxRenderError(RuntimeError):
    pass


_MONTHS_RU = {
    1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель", 5: "Май", 6: "Июнь",
    7: "Июль", 8: "Август", 9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь",
}


# -------------------------
# paths / names
# -------------------------

def _safe_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _month_folder_name(act: Act) -> str:
    m = int(act.act_date.month)
    return f"{m:02d}. {_MONTHS_RU.get(m, str(m))}"


def _project_codes(act: Act) -> list[str]:
    codes: list[str] = []
    for p in act.projects.all().order_by("id"):
        code = (getattr(p, "full_code", "") or "").strip() or str(p).strip()
        if code:
            codes.append(code)
    return codes


# -------------------------
# token replace {{ ... }} (быстро, без зависаний)
# -------------------------

def _normalize_invisible(s: str) -> str:
    return (s or "").replace("\u00A0", " ").replace("\u200b", "")


def _build_token_regex(keys: list[str]) -> re.Pattern:
    keys = sorted(keys, key=len, reverse=True)
    inner = "|".join(re.escape(k) for k in keys)
    return re.compile(r"\{\{\s*(?P<key>" + inner + r")\s*\}\}")


def _replace_tokens_in_paragraph_runs(paragraph, mapping: dict[str, str], token_re: re.Pattern) -> None:
    runs = paragraph.runs
    if not runs:
        return

    full = "".join(r.text for r in runs)
    if "{{" not in full:
        return

    full_norm = _normalize_invisible(full)
    if not token_re.search(full_norm):
        return

    max_ops = 80  # чуть выше, т.к. в некоторых строках много токенов
    ops = 0

    while ops < max_ops:
        full_norm = _normalize_invisible("".join(r.text for r in runs))
        m = token_re.search(full_norm)
        if not m:
            break

        key = m.group("key")
        value = mapping.get(key, "") or ""

        start = m.start()
        end = m.end()

        pos_map = []
        for ri, r in enumerate(runs):
            t = _normalize_invisible(r.text)
            for oi in range(len(t)):
                pos_map.append((ri, oi))

        if len(pos_map) < end:
            break

        (sr, so) = pos_map[start]
        (er, eo) = pos_map[end - 1]

        first = runs[sr]
        last = runs[er]

        prefix = first.text[:so]
        suffix = last.text[eo + 1:] if er != sr else first.text[eo + 1:]

        first.text = prefix + value + suffix

        for j in range(sr + 1, er + 1):
            runs[j].text = ""

        ops += 1


def _replace_in_table(table, mapping: dict[str, str], token_re: re.Pattern) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_tokens_in_paragraph_runs(p, mapping, token_re)
            for inner in cell.tables:
                _replace_in_table(inner, mapping, token_re)


def replace_tokens(document: Document, mapping: dict[str, str]) -> None:
    token_re = _build_token_regex(list(mapping.keys()))

    for p in document.paragraphs:
        _replace_tokens_in_paragraph_runs(p, mapping, token_re)

    for t in document.tables:
        _replace_in_table(t, mapping, token_re)


# -------------------------
# FIX #2: approvals fallback (надёжный)
# -------------------------

_APPROVALS_TOKEN_RE = re.compile(r"\{\{\s*approvals\s*\}\}")


def _force_replace_token_in_paragraph(paragraph, token_re: re.Pattern, value: str) -> bool:
    """
    Делает замену token_re на value даже если токен размазан по runs.
    Возвращает True если что-то заменили.
    """
    runs = paragraph.runs
    if not runs:
        return False

    full = "".join(r.text for r in runs)
    if "{{" not in full:
        return False

    full_norm = _normalize_invisible(full)

    if not token_re.search(full_norm):
        return False

    # Самый стабильный путь: переписать ТОЛЬКО текст runs, не трогая формат runs.
    # Мы:
    # - собираем full исходный (raw)
    # - делаем замену в raw (но ищем в нормализованном)
    # Чтобы не усложнять маппинг, делаем так:
    #   1) Если токен есть в raw напрямую — заменяем raw.replace
    #   2) Иначе (разбит пробелами/невидимыми) — заменяем в norm и потом кладём в первый run,
    #      оставляя стиль первого run. Это не влияет на стили документа в целом (обычная строка).
    if token_re.search(full):
        replaced = token_re.sub(value, full)
        runs[0].text = replaced
        for r in runs[1:]:
            r.text = ""
        return True

    replaced = token_re.sub(value, full_norm)
    runs[0].text = replaced
    for r in runs[1:]:
        r.text = ""
    return True


def force_approvals_everywhere(doc: Document, approvals_value: str) -> None:
    approvals_value = approvals_value or ""

    for p in doc.paragraphs:
        _force_replace_token_in_paragraph(p, _APPROVALS_TOKEN_RE, approvals_value)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _force_replace_token_in_paragraph(p, _APPROVALS_TOKEN_RE, approvals_value)


# -------------------------
# FIX #1: remove trailing ","
# WITHOUT rewriting paragraph (preserve styles)
# -------------------------

_TRAILING_G_COMMA_RE = re.compile(r"(г\.)\s*,\s*$")


def remove_trailing_comma_after_g_dot(paragraph) -> None:
    """
    Если параграф заканчивается на "г.," или "г. ," — удаляем только запятую в последнем run.
    Форматирование НЕ трогаем.
    """
    txt = (paragraph.text or "")
    if not txt.strip():
        return
    if not _TRAILING_G_COMMA_RE.search(txt.strip()):
        return

    # идём по runs с конца, ищем первую запятую, которая стоит "в конце"
    # и удаляем только её (и возможные пробелы после неё).
    for r in reversed(paragraph.runs):
        if not r.text:
            continue
        # уберём хвостовые пробелы/табуляции в этом run (только справа)
        rt = r.text
        # если в этом run вообще нет запятой — дальше
        if "," not in rt:
            continue

        # проверяем: после самой правой запятой в ЭТОМ run только пробелы/переводы
        idx = rt.rfind(",")
        tail = rt[idx + 1:]
        if tail.strip():
            continue  # запятая не хвостовая

        # удаляем запятую и хвостовые пробелы после неё
        r.text = rt[:idx] + tail  # tail тут только пробельный, оставим как было
        return


def fix_exec_scheme_comma_in_document(doc: Document) -> None:
    for p in doc.paragraphs:
        remove_trailing_comma_after_g_dot(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    remove_trailing_comma_after_g_dot(p)


# -------------------------
# Appendices table
# -------------------------

def _cell_text(cell) -> str:
    return "\n".join((p.text or "") for p in cell.paragraphs).strip()


def _find_appendix_table(doc: Document):
    for t in doc.tables:
        if not t.rows:
            continue
        first_row_texts = [(_cell_text(c) or "").strip() for c in t.rows[0].cells]
        if any(txt == "Приложения:" for txt in first_row_texts):
            return t
    return None


def _normalize_line_text(s: str) -> str:
    s = (s or "").strip()
    while s.endswith(","):
        s = s[:-1].rstrip()
    s = re.sub(r"(\d{2}\.\d{2}\.\d{4})(?!г\.)", r"\1г.", s)
    return s


def _copy_paragraph_format(src_p, dst_p) -> None:
    try:
        dst_p.style = src_p.style
    except Exception:
        pass
    try:
        dst_p.paragraph_format.left_indent = src_p.paragraph_format.left_indent
        dst_p.paragraph_format.first_line_indent = src_p.paragraph_format.first_line_indent
        dst_p.paragraph_format.space_before = src_p.paragraph_format.space_before
        dst_p.paragraph_format.space_after = src_p.paragraph_format.space_after
        dst_p.paragraph_format.line_spacing = src_p.paragraph_format.line_spacing
        dst_p.paragraph_format.alignment = src_p.paragraph_format.alignment
    except Exception:
        pass


def _copy_run_style(src_run, dst_run) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline

    sf = src_run.font
    df = dst_run.font
    df.name = sf.name
    df.size = sf.size
    if sf.color and sf.color.rgb:
        df.color.rgb = sf.color.rgb


def _set_cell_text_keep_format(cell, text: str, template_cell=None) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

    for r in p.runs:
        r.text = ""

    tpl_run = None
    if template_cell and template_cell.paragraphs:
        tp = template_cell.paragraphs[0]
        _copy_paragraph_format(tp, p)
        if tp.runs:
            tpl_run = tp.runs[0]

    new_run = p.add_run(text or "")

    if tpl_run is not None:
        _copy_run_style(tpl_run, new_run)

    # насильно Times New Roman 11 курсив (как ты просил — со 2-й строки вниз)
    new_run.font.name = "Times New Roman"
    new_run.font.size = Pt(11)
    new_run.italic = True


def _delete_row(table, row_idx: int) -> None:
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    tbl.remove(tr)


def fill_appendix_table(doc: Document, act: Act) -> None:
    table = _find_appendix_table(doc)
    if table is None:
        raise DocxRenderError("Не найдена таблица, где первая строка содержит 'Приложения:'.")

    if len(table.columns) < 2:
        raise DocxRenderError("Таблица 'Приложения' должна иметь минимум 2 колонки (№ и документ).")

    lines = list(act.appendix_lines.all().order_by("position", "id"))
    if not lines:
        return

    template_row_idx = 1
    if len(table.rows) <= template_row_idx:
        table.add_row()

    template_row = table.rows[template_row_idx]
    t_no = template_row.cells[0]
    t_doc = template_row.cells[1]

    needed_rows = 1 + len(lines)
    while len(table.rows) < needed_rows:
        table.add_row()
    while len(table.rows) > needed_rows:
        _delete_row(table, len(table.rows) - 1)

    for i, line in enumerate(lines, start=1):
        row = table.rows[i]
        _set_cell_text_keep_format(row.cells[0], f"{i}.", template_cell=t_no)
        _set_cell_text_keep_format(row.cells[1], _normalize_line_text(str(line.label or "")), template_cell=t_doc)


# -------------------------
# main
# -------------------------

from pathlib import Path
from typing import Optional, List

from django.conf import settings
from docx import Document

# ... (остальные импорты и код файла остаются как у тебя)

def get_act_docx_paths(act: Act) -> list[Path]:
    """
    Ожидаемые пути, куда должен быть сохранён docx, без сохранения.
    Формат:
      ACTS_DIR / year / month_folder / project_full_code / "Акт №... от dd.mm.yyyy.docx"
    """
    year_dir = Path(settings.ACTS_DIR) / f"{act.act_date.year}"
    month_dir = year_dir / _month_folder_name(act)

    project_codes = _project_codes(act) or ["Без проекта"]
    file_name = _safe_filename(f"Акт №{act.number} от {act.act_date:%d.%m.%Y}") + ".docx"

    return [(month_dir / _safe_filename(code) / file_name) for code in project_codes]


def generate_act_docx(act: Act, *, template_path: Optional[Path] = None) -> list[Path]:
    """
    Генерирует docx и сохраняет его в вычисленные пути (перезаписывает если существует).
    Возвращает список путей.
    """
    template_path = template_path or (Path(settings.DOCX_TEMPLATES_DIR) / "act_template.docx")
    if not template_path.exists():
        raise DocxRenderError(f"Не найден DOCX-шаблон: {template_path}")

    # пересобираем приложения каждый раз перед генерацией
    AppendixBuilder(act).rebuild()

    # грузим шаблон
    doc = Document(str(template_path))

    # собираем контекст
    ctx = build_act_docx_context(act)

    # ✅ ВАЖНО: у тебя в проекте это называется replace_tokens, а не _fill_tokens
    replace_tokens(doc, ctx)

    # приложения внутри docx
    fill_appendix_table(doc, act)

    # твои фиксы
    fix_exec_scheme_comma_in_document(doc)
    force_approvals_everywhere(doc, ctx.get("approvals", ""))

    # путь с year/ в начале
    paths = get_act_docx_paths(act)

    saved: list[Path] = []
    for out_path in paths:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        saved.append(out_path)

    return saved

