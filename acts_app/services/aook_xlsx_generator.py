from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from copy import copy, deepcopy
from pathlib import Path
from typing import Any

from django.conf import settings
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from openpyxl.cell.cell import MergedCell
from openpyxl.styles import Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.workbook.workbook import Workbook
from openpyxl.worksheet.worksheet import Worksheet

from acts_app.models import Aook, AookManualSourceAct, AookProtocolItem, AookSourceAct
from acts_app.services.act_docx_generator import (
    _build_token_regex,
    _replace_tokens_in_paragraph_runs,
    _month_folder_name,
    _safe_filename as _safe_act_filename,
    replace_tokens,
)
from acts_app.services.act_docx_context import build_act_docx_context
from acts_app.services.date_format import fmt_date_g
from documents_app.utils.pdf_utils import (
    _get_libreoffice_executable,
    _get_libreoffice_profile_uri,
    convert_docx_to_pdf,
)


class AookXlsxRenderError(RuntimeError):
    pass


class AookRenderError(RuntimeError):
    pass


_TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")
_ATTACHMENT_TOKEN_RE = re.compile(r"\{\{\s*attachment\s*\}\}")
_ATTACHMENT_NUMBER_RE = re.compile(r"^\s*\d+\.\s*")
_STRAIGHT_QUOTED_TEXT_RE = re.compile(r'"([^"\r\n]+)"')
_THIN_BOTTOM_SIDE = Side(style="thin", color="000000")
_MONTHS_RU_GENITIVE = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}


def _safe_filename(value: str) -> str:
    value = (value or "").strip()
    value = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value[:160] or "aook"


def _date_parts(prefix: str, value) -> dict[str, str]:
    if not value:
        return {
            f"{prefix}_day": "",
            f"{prefix}_month": "",
            f"{prefix}_year": "",
        }
    return {
        f"{prefix}_day": f"{value.day:02d}",
        f"{prefix}_month": _MONTHS_RU_GENITIVE.get(int(value.month), ""),
        f"{prefix}_year": f"{value.year % 100:02d}",
    }


def _registry_date_parts(prefix: str, value) -> dict[str, str]:
    if not value:
        return {
            f"{prefix}_day": "",
            f"{prefix}_month": "",
            f"{prefix}_year": "",
        }
    return {
        f"{prefix}_day": f"{value.day:02d}",
        f"{prefix}_month": f"{value.month:02d}",
        f"{prefix}_year": f"{value.year % 100:02d}",
    }


def _format_date_short(value) -> str:
    return value.strftime("%d.%m.%y") if value else ""


def _format_date_full(value) -> str:
    return value.strftime("%d.%m.%Y") if value else ""


def _typograph_quotes(value: Any) -> Any:
    if not isinstance(value, str) or '"' not in value:
        return value
    return _STRAIGHT_QUOTED_TEXT_RE.sub(r"«\1»", value)


def _typograph_mapping_quotes(mapping: dict[str, Any]) -> dict[str, Any]:
    return {key: _typograph_quotes(value) for key, value in mapping.items()}


def _format_act_registry_date(act) -> str:
    return _format_date_full(getattr(act, "act_date", None))


def _format_manual_source_act_date(item: AookManualSourceAct) -> str:
    return _format_date_full(item.act_date)


def _format_protocol_date(item: AookProtocolItem) -> str:
    return _format_date_full(item.document_date)


def _build_attachments_text(aook: Aook) -> str:
    return "\n".join(
        _build_attachment_lines(aook)
    )


def _build_attachment_lines(aook: Aook) -> list[str]:
    return [
        f"1. исполнительная схема №{aook.number} от {fmt_date_g(aook.act_date)}",
        f"2. реестр актов №{aook.aosr_registry_number or f'П-3.{aook.number}'} от {fmt_date_g(aook.act_date)}",
        f"3. реестр протоколов №{aook.protocols_registry_number or f'П-6.{aook.number}'} от {fmt_date_g(aook.act_date)}",
    ]


def _attachment_line_text(line: str) -> str:
    return _ATTACHMENT_NUMBER_RE.sub("", line or "").strip()


def _with_bottom_border(border: Border | None) -> Border:
    border = copy(border) if border else Border()
    return Border(
        left=copy(border.left),
        right=copy(border.right),
        top=copy(border.top),
        bottom=_THIN_BOTTOM_SIDE,
        diagonal=copy(border.diagonal),
        diagonal_direction=border.diagonal_direction,
        diagonalUp=border.diagonalUp,
        diagonalDown=border.diagonalDown,
        outline=border.outline,
        vertical=copy(border.vertical),
        horizontal=copy(border.horizontal),
        start=copy(border.start),
        end=copy(border.end),
    )


def _apply_excel_row_bottom_border(worksheet: Worksheet, row_idx: int) -> None:
    for col_idx in range(1, worksheet.max_column + 1):
        cell = worksheet.cell(row=row_idx, column=col_idx)
        if isinstance(cell, MergedCell):
            continue
        cell.border = _with_bottom_border(cell.border)


def _replace_attachment_placeholder(workbook: Workbook, aook: Aook) -> dict[str, dict[int, set[int]]]:
    touched: dict[str, dict[int, set[int]]] = {}
    lines = _build_attachment_lines(aook)
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows():
            for cell in row:
                if isinstance(cell, MergedCell):
                    continue
                if not (isinstance(cell.value, str) and "{{attachment}}" in cell.value):
                    continue

                start_row = cell.row
                column = cell.column
                number_column = column - 1 if column > 1 else None
                has_number_column = False
                if number_column:
                    number_cell = worksheet.cell(row=start_row, column=number_column)
                    has_number_column = not isinstance(number_cell, MergedCell) and number_cell.value not in (None, "")

                for offset, line in enumerate(lines):
                    target = worksheet.cell(row=start_row + offset, column=column)
                    if isinstance(target, MergedCell):
                        continue
                    if offset:
                        _copy_row_style(
                            worksheet,
                            source_row=start_row,
                            target_row=start_row + offset,
                            copy_values=False,
                        )
                    if has_number_column and number_column:
                        number_cell = worksheet.cell(row=start_row + offset, column=number_column)
                        if not isinstance(number_cell, MergedCell):
                            number_cell.value = f"{offset + 1}."
                        target.value = _attachment_line_text(line)
                    else:
                        target.value = line
                    _apply_excel_row_bottom_border(worksheet, start_row + offset)
                    for attachment_cell in (target, worksheet.cell(row=start_row + offset, column=number_column) if number_column else None):
                        if attachment_cell is None or isinstance(attachment_cell, MergedCell):
                            continue
                        alignment = copy(attachment_cell.alignment) if attachment_cell.alignment else Alignment()
                        alignment.wrap_text = False
                        alignment.vertical = alignment.vertical or "center"
                        attachment_cell.alignment = alignment
                    worksheet.row_dimensions[start_row + offset].height = _default_row_height(worksheet)
                    touched.setdefault(worksheet.title, {}).setdefault(start_row + offset, set()).add(column)
                    if number_column:
                        touched[worksheet.title][start_row + offset].add(number_column)
                return touched
    return touched


def _docx_cell_text(cell) -> str:
    return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()


def _docx_cell_has_attachment_token(cell) -> bool:
    return any(_ATTACHMENT_TOKEN_RE.search(paragraph.text or "") for paragraph in cell.paragraphs)


def _set_docx_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()

    first_paragraph = cell.paragraphs[0]
    if first_paragraph.runs:
        first_paragraph.runs[0].text = text
        for run in first_paragraph.runs[1:]:
            run.text = ""
    else:
        first_paragraph.add_run(text)

    for paragraph in cell.paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def _set_docx_cell_bottom_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")


def _set_docx_row_bottom_border(row) -> None:
    for cell in row.cells:
        _set_docx_cell_bottom_border(cell)


def _iter_docx_tables(document_or_cell):
    for table in document_or_cell.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_docx_tables(cell)


def _insert_docx_row_after(table, row_idx: int):
    new_tr = deepcopy(table.rows[row_idx]._tr)
    table.rows[row_idx]._tr.addnext(new_tr)
    return table.rows[row_idx + 1]


def _replace_docx_attachment_placeholder(document: Document, aook: Aook) -> None:
    lines = _build_attachment_lines(aook)
    for table in _iter_docx_tables(document):
        for row_idx, row in enumerate(table.rows):
            placeholder_cell_idx = None
            for cell_idx, cell in enumerate(row.cells):
                if _docx_cell_has_attachment_token(cell):
                    placeholder_cell_idx = cell_idx
                    break
            if placeholder_cell_idx is None:
                continue

            for offset in range(1, len(lines)):
                _insert_docx_row_after(table, row_idx + offset - 1)

            for offset, line in enumerate(lines):
                target_row = table.rows[row_idx + offset]
                target_cell = target_row.cells[placeholder_cell_idx]
                previous_cell = target_row.cells[placeholder_cell_idx - 1] if placeholder_cell_idx > 0 else None
                previous_text = _docx_cell_text(previous_cell) if previous_cell is not None else ""
                has_number_cell = previous_cell is not None and previous_text in {"", "1", "1."}

                if has_number_cell:
                    _set_docx_cell_text(previous_cell, f"{offset + 1}.")
                    _set_docx_cell_text(target_cell, _attachment_line_text(line))
                else:
                    _set_docx_cell_text(target_cell, f"{offset + 1}. {_attachment_line_text(line)}")
                    if offset and previous_cell is not None:
                        _set_docx_cell_text(previous_cell, "")

                _set_docx_row_bottom_border(target_row)
            return


def _replace_tokens_in_docx_row(row, mapping: dict[str, str]) -> None:
    if not mapping:
        return
    token_re = _build_token_regex(list(mapping.keys()))
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            _replace_tokens_in_paragraph_runs(paragraph, mapping, token_re)
        for table in cell.tables:
            for nested_row in table.rows:
                _replace_tokens_in_docx_row(nested_row, mapping)


def _remove_docx_row(row) -> None:
    row._tr.getparent().remove(row._tr)


def _row_contains_any_key(row, keys: set[str]) -> bool:
    row_text = "\n".join(_docx_cell_text(cell) for cell in row.cells)
    return any(match.group(1).strip() in keys for match in _TOKEN_RE.finditer(row_text))


def _docx_row_is_empty(row) -> bool:
    return not "".join(_docx_cell_text(cell) for cell in row.cells).strip()


def _fill_docx_registry_rows(document: Document, rows: list[dict[str, Any]], row_keys: set[str]) -> None:
    for table in _iter_docx_tables(document):
        template_row_indexes = [
            row_idx
            for row_idx, row in enumerate(table.rows)
            if _row_contains_any_key(row, row_keys)
        ]
        if not template_row_indexes:
            continue

        blank_row_indexes = []
        row_idx = template_row_indexes[-1] + 1
        while row_idx < len(table.rows) and _docx_row_is_empty(table.rows[row_idx]):
            blank_row_indexes.append(row_idx)
            row_idx += 1

        for row_idx in reversed(blank_row_indexes):
            _remove_docx_row(table.rows[row_idx])

        if not rows:
            for row_idx in reversed(template_row_indexes):
                _remove_docx_row(table.rows[row_idx])
            return

        while len(template_row_indexes) < len(rows):
            insert_after_idx = template_row_indexes[-1]
            _insert_docx_row_after(table, insert_after_idx)
            template_row_indexes.append(insert_after_idx + 1)

        for offset, row_mapping in enumerate(rows):
            target_row = table.rows[template_row_indexes[offset]]
            _replace_tokens_in_docx_row(
                target_row,
                {key: str(_typograph_quotes(value) or "") for key, value in row_mapping.items()},
            )

        for row_idx in reversed(template_row_indexes[len(rows):]):
            _remove_docx_row(table.rows[row_idx])
        return

    raise AookRenderError(f"В DOCX-шаблоне реестра АООК не найдена строка с плейсхолдерами: {', '.join(sorted(row_keys))}")


def _copy_row_style(worksheet: Worksheet, *, source_row: int, target_row: int, copy_values: bool = True) -> None:
    max_col = worksheet.max_column
    for col_idx in range(1, max_col + 1):
        source = worksheet.cell(row=source_row, column=col_idx)
        target = worksheet.cell(row=target_row, column=col_idx)
        if copy_values and not isinstance(target, MergedCell):
            target.value = source.value
        if source.has_style:
            target._style = copy(source._style)
        if source.number_format:
            target.number_format = source.number_format
        if source.font:
            target.font = copy(source.font)
        if source.fill:
            target.fill = copy(source.fill)
        if source.border:
            target.border = copy(source.border)
        if source.alignment:
            target.alignment = copy(source.alignment)
        if source.protection:
            target.protection = copy(source.protection)
    worksheet.row_dimensions[target_row].height = worksheet.row_dimensions[source_row].height


def _replicate_merged_ranges_for_inserted_rows(worksheet: Worksheet, *, template_row: int, rows_count: int) -> None:
    ranges = list(worksheet.merged_cells.ranges)
    for merged_range in ranges:
        if merged_range.min_row != template_row or merged_range.max_row != template_row:
            continue
        for row_idx in range(template_row + 1, template_row + rows_count):
            worksheet.merge_cells(
                start_row=row_idx,
                start_column=merged_range.min_col,
                end_row=row_idx,
                end_column=merged_range.max_col,
            )


def _replace_placeholders(worksheet: Worksheet, mapping: dict[str, Any]) -> dict[int, set[int]]:
    touched: dict[int, set[int]] = {}

    def repl(match: re.Match) -> str:
        key = match.group(1)
        return str(mapping.get(key, "") or "")

    for row in worksheet.iter_rows():
        for cell in row:
            if isinstance(cell, MergedCell):
                continue
            if isinstance(cell.value, str) and "{{" in cell.value:
                cell.value = _TOKEN_RE.sub(repl, cell.value)
                touched.setdefault(cell.row, set()).add(cell.column)
    return touched


def _write_registry_rows(
    worksheet: Worksheet,
    *,
    template_row: int,
    rows: list[dict[str, Any]],
    mapping_prefix: str,
) -> set[int]:
    touched: set[int] = set()
    has_rows = bool(rows)
    rows = rows or [{}]
    additional = max(0, len(rows) - 1)
    if additional:
        worksheet.insert_rows(template_row + 1, amount=additional)
        _replicate_merged_ranges_for_inserted_rows(worksheet, template_row=template_row, rows_count=len(rows))

    for offset, row_mapping in enumerate(rows):
        row_idx = template_row + offset
        if offset:
            _copy_row_style(worksheet, source_row=template_row, target_row=row_idx)
        first_cell = worksheet.cell(row=row_idx, column=1)
        if not isinstance(first_cell, MergedCell):
            first_cell.value = offset + 1 if has_rows else ""
        mapping = {
            f"{mapping_prefix}_pp": str(offset + 1),
            **row_mapping,
        }
        _replace_placeholders_in_row(worksheet, row_idx=row_idx, mapping=mapping)
        touched.add(row_idx)
    return touched


def _replace_placeholders_in_row(worksheet: Worksheet, *, row_idx: int, mapping: dict[str, Any]) -> None:
    def repl(match: re.Match) -> str:
        key = match.group(1)
        return str(mapping.get(key, "") or "")

    for cell in worksheet[row_idx]:
        if isinstance(cell, MergedCell):
            continue
        if isinstance(cell.value, str) and "{{" in cell.value:
            cell.value = _TOKEN_RE.sub(repl, cell.value)


def _cell_text_width(worksheet: Worksheet, cell) -> float:
    for merged_range in worksheet.merged_cells.ranges:
        if cell.coordinate in merged_range:
            width = 0.0
            for col_idx in range(merged_range.min_col, merged_range.max_col + 1):
                letter = get_column_letter(col_idx)
                width += float(worksheet.column_dimensions[letter].width or 10)
            return max(width, 10.0)

    letter = get_column_letter(cell.column)
    return float(worksheet.column_dimensions[letter].width or 10)


def _estimated_text_lines(text: str, width: float) -> int:
    chars_per_line = max(1, int(width * 1.25))
    total = 0
    for part in str(text or "").splitlines() or [""]:
        part = part.strip()
        total += max(1, (len(part) + chars_per_line - 1) // chars_per_line)
    return max(1, total)


def _default_row_height(worksheet: Worksheet) -> float:
    return float(worksheet.sheet_format.defaultRowHeight or 15)


def _fit_rows_to_text(
    worksheet: Worksheet,
    row_numbers: set[int],
    columns_by_row: dict[int, set[int]] | None = None,
    *,
    min_height: float | None = None,
) -> None:
    min_height = min_height if min_height is not None else _default_row_height(worksheet)
    for row_number in sorted(row_numbers):
        row = tuple(worksheet[row_number])
        selected_columns = columns_by_row.get(row_number) if columns_by_row else None
        max_lines = 1
        for cell in row:
            if isinstance(cell, MergedCell) or cell.value in (None, ""):
                continue
            if selected_columns is not None and cell.column not in selected_columns:
                continue

            alignment = copy(cell.alignment) if cell.alignment else Alignment()
            alignment.wrap_text = True
            alignment.vertical = alignment.vertical or "center"
            cell.alignment = alignment

            if isinstance(cell.value, str):
                max_lines = max(max_lines, _estimated_text_lines(cell.value, _cell_text_width(worksheet, cell)))

        row_index = row[0].row if row else 1
        if max_lines <= 1:
            worksheet.row_dimensions[row_index].height = min_height
        else:
            worksheet.row_dimensions[row_index].height = max(min_height, min(300, max_lines * min_height))


def _merge_touched_cells(target: dict[int, set[int]], source: dict[int, set[int]]) -> None:
    for row_number, columns in source.items():
        target.setdefault(row_number, set()).update(columns)


def _normalize_sheet_geometry(worksheet: Worksheet, *, max_column: int = 36) -> None:
    if worksheet.max_column > max_column:
        worksheet.delete_cols(max_column + 1, worksheet.max_column - max_column)
    for column_idx in range(1, max_column + 1):
        worksheet.column_dimensions[get_column_letter(column_idx)].width = 2.77
    worksheet.print_area = f"A1:{get_column_letter(max_column)}{worksheet.max_row}"


def _resave_xlsx_with_libreoffice(source_path: Path, tmp_dir: Path) -> Path:
    output_dir = tmp_dir / "lo_measured"
    output_dir.mkdir(parents=True, exist_ok=True)
    executable = _get_libreoffice_executable()
    args = [
        executable,
        f"-env:UserInstallation={_get_libreoffice_profile_uri()}",
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        "xlsx",
        "--outdir",
        str(output_dir),
        str(source_path),
    ]
    proc = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=False, timeout=180)
    if proc.returncode != 0:
        raise AookXlsxRenderError(
            "LibreOffice не смог пересохранить измерительный XLSX для расчета высот строк. "
            f"stdout: {proc.stdout.strip() or '-'}; stderr: {proc.stderr.strip() or '-'}"
        )

    measured_path = output_dir / source_path.name
    if not measured_path.exists():
        candidates = list(output_dir.glob("*.xlsx"))
        if candidates:
            measured_path = candidates[0]
    if not measured_path.exists():
        raise AookXlsxRenderError("LibreOffice не создал измерительный XLSX для расчета высот строк.")
    return measured_path


def _measurement_rows(
    rows_by_sheet: dict[str, set[int]],
    cells_by_sheet: dict[str, dict[int, set[int]]],
) -> dict[str, set[int]]:
    result: dict[str, set[int]] = {}
    for sheet_name, rows in rows_by_sheet.items():
        result.setdefault(sheet_name, set()).update(rows)
    for sheet_name, cells_by_row in cells_by_sheet.items():
        result.setdefault(sheet_name, set()).update(cells_by_row)
    return result


def _prepare_measurement_workbook(
    workbook: Workbook,
    *,
    rows_by_sheet: dict[str, set[int]],
    full_row_sheets: dict[str, set[int]],
    cells_by_sheet: dict[str, dict[int, set[int]]],
) -> None:
    for worksheet in workbook.worksheets:
        sheet_rows = rows_by_sheet.get(worksheet.title, set())
        full_rows = full_row_sheets.get(worksheet.title, set())
        cells_by_row = cells_by_sheet.get(worksheet.title, {})
        for row_number in sorted(sheet_rows):
            selected_columns = cells_by_row.get(row_number)
            measure_full_row = row_number in full_rows or not selected_columns
            for cell in worksheet[row_number]:
                if isinstance(cell, MergedCell):
                    continue
                if not measure_full_row and cell.column not in selected_columns:
                    cell.value = None
                    continue
                alignment = copy(cell.alignment) if cell.alignment else Alignment()
                alignment.wrap_text = True
                alignment.vertical = alignment.vertical or "center"
                cell.alignment = alignment
            worksheet.row_dimensions[row_number].height = None


def _apply_measured_alignment(
    workbook: Workbook,
    *,
    full_row_sheets: dict[str, set[int]],
    cells_by_sheet: dict[str, dict[int, set[int]]],
) -> None:
    rows_by_sheet = _measurement_rows(full_row_sheets, cells_by_sheet)
    for worksheet in workbook.worksheets:
        full_rows = full_row_sheets.get(worksheet.title, set())
        cells_by_row = cells_by_sheet.get(worksheet.title, {})
        for row_number in sorted(rows_by_sheet.get(worksheet.title, set())):
            selected_columns = cells_by_row.get(row_number)
            measure_full_row = row_number in full_rows or not selected_columns
            for cell in worksheet[row_number]:
                if isinstance(cell, MergedCell) or cell.value in (None, ""):
                    continue
                if not measure_full_row and cell.column not in selected_columns:
                    continue
                alignment = copy(cell.alignment) if cell.alignment else Alignment()
                alignment.wrap_text = True
                alignment.vertical = alignment.vertical or "center"
                cell.alignment = alignment


def _apply_measured_row_heights(
    workbook: Workbook,
    *,
    measured_path: Path,
    rows_by_sheet: dict[str, set[int]],
) -> None:
    measured_workbook = load_workbook(measured_path)
    try:
        for worksheet in workbook.worksheets:
            if worksheet.title not in measured_workbook.sheetnames:
                continue
            measured_sheet = measured_workbook[worksheet.title]
            for row_number in rows_by_sheet.get(worksheet.title, set()):
                measured_height = measured_sheet.row_dimensions[row_number].height
                worksheet.row_dimensions[row_number].height = measured_height or _default_row_height(worksheet)
    finally:
        close = getattr(measured_workbook, "close", None)
        if close:
            close()


def _normalize_sheet_name(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip().lower())


def _resolve_sheet(workbook: Workbook, *names: str) -> Worksheet:
    candidates = {_normalize_sheet_name(name) for name in names}
    for worksheet in workbook.worksheets:
        if _normalize_sheet_name(worksheet.title) in candidates:
            return worksheet
    available = ", ".join(workbook.sheetnames)
    expected = ", ".join(names)
    raise AookXlsxRenderError(f"В шаблоне АООК не найден лист: {expected}. Доступные листы: {available}")


def _convert_xls_template_to_xlsx(template_path: Path, tmp_dir: Path) -> Path:
    if template_path.suffix.lower() == ".xlsx":
        return template_path

    if template_path.suffix.lower() != ".xls":
        raise AookXlsxRenderError(f"Неизвестный формат шаблона АООК: {template_path}")

    executable = _get_libreoffice_executable()
    args = [
        executable,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        "xlsx",
        "--outdir",
        str(tmp_dir),
        str(template_path),
    ]
    proc = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=False, timeout=180)
    if proc.returncode != 0:
        raise AookXlsxRenderError(
            "LibreOffice не смог конвертировать XLS-шаблон АООК в XLSX. "
            f"stdout: {proc.stdout.strip() or '-'}; stderr: {proc.stderr.strip() or '-'}"
        )

    converted = tmp_dir / f"{template_path.stem}.xlsx"
    if not converted.exists():
        candidates = list(tmp_dir.glob("*.xlsx"))
        if candidates:
            converted = candidates[0]
    if not converted.exists():
        raise AookXlsxRenderError(f"LibreOffice не создал XLSX из шаблона: {template_path}")
    return converted


def _source_act_rows(aook: Aook) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    items = (
        aook.source_act_items.select_related("act")
        .order_by("position", "id")
    )
    for position, item in enumerate(items, start=1):
        act = item.act
        rows.append(
            {
                "p_n": position,
                "aosr": "Акт освидетельствования скрытых работ",
                "registry_act_job_name": "Акт освидетельствования скрытых работ",
                "registry_act_number": act.number,
                "registry_act_date": _format_act_registry_date(act),
                "contractor_rep_org_short": _contractor_org_short(act),
            }
        )
    manual_items = (
        aook.manual_source_act_items
        .order_by("position", "id")
    )
    for position, item in enumerate(manual_items, start=len(rows) + 1):
        rows.append(
            {
                "p_n": position,
                "aosr": "Акт освидетельствования скрытых работ",
                "registry_act_job_name": "Акт освидетельствования скрытых работ",
                "registry_act_number": item.act_number,
                "registry_act_date": _format_manual_source_act_date(item),
                "contractor_rep_org_short": item.organization_name,
            }
        )
    return rows


def _protocol_rows(aook: Aook) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for position, item in enumerate(aook.protocol_items.order_by("position", "id"), start=1):
        rows.append(
            {
                "p_n": position,
                "registry_protocol_document_name": item.document_name,
                "registry_protocol_number": item.document_number,
                "registry_protocol_date": _format_protocol_date(item),
                "registry_protocol_org_name": item.organization_name,
            }
        )
    return rows


def _contractor_org_short(act) -> str:
    party = (
        act.parties.select_related("organization")
        .filter(role="CONTRACTOR_REP", is_enabled=True)
        .order_by("position", "id")
        .first()
    )
    org = party.organization if party and party.organization_id else None
    return (getattr(org, "short_name", "") or getattr(org, "full_name", "") or "").strip()


def build_aook_xlsx_context(aook: Aook) -> dict[str, Any]:
    source_item: AookSourceAct | None = (
        aook.source_act_items.select_related("act")
        .order_by("position", "id")
        .first()
    )
    if source_item is None:
        raise AookXlsxRenderError("Для АООК не выбраны исходные АОСР.")

    mapping = build_act_docx_context(source_item.act)
    mapping.update(
        {
            "aook_number": aook.number,
            "job_name": aook.work_name,
            "copies_count": aook.copies_count,
            "work_norms_text": aook.work_norms_text,
            "aosr_registry_number": aook.aosr_registry_number or f"П-3.{aook.number}",
            "aosr_registry_date": fmt_date_g(aook.act_date),
            "protocols_registry_number": aook.protocols_registry_number or f"П-6.{aook.number}",
            "protocols_registry_date": fmt_date_g(aook.act_date),
            "sheme_number": aook.number,
            "scheme_date": fmt_date_g(aook.act_date),
            "attachment": _build_attachments_text(aook),
        }
    )
    mapping.update(_date_parts("aook", aook.act_date))
    mapping.update(_date_parts("aook_job_start", aook.work_start_date))
    mapping.update(_date_parts("aook_job_end", aook.work_end_date))
    mapping.update(
        {
            "registry_contractor_org_full": mapping.get("contractor_rep_org_full", ""),
            "registry_signer_position": mapping.get("contractor_rep_position", ""),
            "registry_signer_fio": mapping.get("contractor_rep_fio", ""),
        }
    )
    return _typograph_mapping_quotes(mapping)


def _apply_aook_registry_date_context(mapping: dict[str, Any], aook: Aook) -> dict[str, Any]:
    registry_mapping = dict(mapping)
    registry_mapping.update(
        {
            "aosr_registry_date": _format_date_short(aook.act_date),
            "protocols_registry_date": _format_date_short(aook.act_date),
        }
    )
    registry_mapping.update(_registry_date_parts("aook", aook.act_date))
    registry_mapping.update(_registry_date_parts("aook_job_start", aook.work_start_date))
    registry_mapping.update(_registry_date_parts("aook_job_end", aook.work_end_date))
    return registry_mapping


def render_aook_xlsx(*, aook: Aook, output_path: Path, template_path: Path | None = None) -> Path:
    template_path = template_path or (Path(settings.XLSX_TEMPLATES_DIR) / "aook_template.xls")
    if not template_path.exists():
        raise AookXlsxRenderError(f"Шаблон АООК не найден: {template_path}")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="doc_helper_aook_") as tmp:
        tmp_dir = Path(tmp)
        workbook_template = _convert_xls_template_to_xlsx(template_path, tmp_dir)
        workbook = load_workbook(workbook_template)

        mapping = build_aook_xlsx_context(aook)
        for worksheet in workbook.worksheets:
            worksheet.sheet_state = "visible"

        fit_rows_by_sheet: dict[str, set[int]] = {}
        fit_cells_by_sheet: dict[str, dict[int, set[int]]] = {}
        for sheet_name, cells_by_row in _replace_attachment_placeholder(workbook, aook).items():
            _merge_touched_cells(fit_cells_by_sheet.setdefault(sheet_name, {}), cells_by_row)

        act_registry_sheet = _resolve_sheet(workbook, "реестр аоср", "реестр актов")
        fit_rows_by_sheet.setdefault(act_registry_sheet.title, set()).update(_write_registry_rows(
            act_registry_sheet,
            template_row=6,
            rows=_source_act_rows(aook),
            mapping_prefix="registry_act",
        ))
        protocol_registry_sheet = _resolve_sheet(workbook, "реестр протоколов")
        fit_rows_by_sheet.setdefault(protocol_registry_sheet.title, set()).update(_write_registry_rows(
            protocol_registry_sheet,
            template_row=6,
            rows=_protocol_rows(aook),
            mapping_prefix="registry_protocol",
        ))

        for worksheet in workbook.worksheets:
            _merge_touched_cells(
                fit_cells_by_sheet.setdefault(worksheet.title, {}),
                _replace_placeholders(worksheet, mapping),
            )
            _normalize_sheet_geometry(worksheet)

        rows_to_measure = _measurement_rows(fit_rows_by_sheet, fit_cells_by_sheet)
        _apply_measured_alignment(
            workbook,
            full_row_sheets=fit_rows_by_sheet,
            cells_by_sheet=fit_cells_by_sheet,
        )
        draft_path = tmp_dir / "aook_layout_draft.xlsx"
        measurement_input_path = tmp_dir / "aook_layout_measure.xlsx"
        workbook.save(draft_path)
        measurement_workbook = load_workbook(draft_path)
        try:
            _prepare_measurement_workbook(
                measurement_workbook,
                rows_by_sheet=rows_to_measure,
                full_row_sheets=fit_rows_by_sheet,
                cells_by_sheet=fit_cells_by_sheet,
            )
            measurement_workbook.save(measurement_input_path)
        finally:
            close = getattr(measurement_workbook, "close", None)
            if close:
                close()
        measured_path = _resave_xlsx_with_libreoffice(measurement_input_path, tmp_dir)
        _apply_measured_row_heights(workbook, measured_path=measured_path, rows_by_sheet=rows_to_measure)

        workbook.save(output_path)

    return output_path


def render_aook_docx(*, aook: Aook, output_path: Path, template_path: Path | None = None) -> Path:
    template_path = template_path or (Path(settings.DOCX_TEMPLATES_DIR) / "aook_docx_template.docx")
    if not template_path.exists():
        raise AookXlsxRenderError(f"DOCX-шаблон АООК не найден: {template_path}")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document(str(template_path))
    mapping = build_aook_xlsx_context(aook)
    _replace_docx_attachment_placeholder(document, aook)
    replace_tokens(document, {key: str(value or "") for key, value in mapping.items()})
    document.save(str(output_path))
    return output_path


def render_aook_registry_docx(
    *,
    aook: Aook,
    output_path: Path,
    registry_type: str,
    template_path: Path | None = None,
) -> Path:
    if registry_type == "acts":
        template_name = "aook_acts_registry.docx"
        rows = _source_act_rows(aook)
        row_keys = {"p_n", "aosr", "registry_act_job_name", "registry_act_number", "registry_act_date", "contractor_rep_org_short"}
    elif registry_type == "protocols":
        template_name = "aook_protocols_registry.docx"
        rows = _protocol_rows(aook)
        row_keys = {"p_n", "registry_protocol_document_name", "registry_protocol_number", "registry_protocol_date", "registry_protocol_org_name"}
    else:
        raise AookRenderError(f"Неизвестный тип реестра АООК: {registry_type}")

    template_path = template_path or (Path(settings.DOCX_TEMPLATES_DIR) / template_name)
    if not template_path.exists():
        raise AookRenderError(f"DOCX-шаблон реестра АООК не найден: {template_path}")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document(str(template_path))
    mapping = _apply_aook_registry_date_context(build_aook_xlsx_context(aook), aook)
    common_mapping = {key: value for key, value in mapping.items() if key not in row_keys}
    replace_tokens(document, {key: str(value or "") for key, value in common_mapping.items()})
    _fill_docx_registry_rows(document, rows, row_keys)
    document.save(str(output_path))
    return output_path


def get_aook_generated_paths(aook: Aook) -> dict[str, Path]:
    file_stem = _safe_filename(f"АООК №{aook.number} от {aook.act_date:%d.%m.%Y}")
    year_dir = Path(settings.ACTS_DIR) / f"{aook.act_date.year}"
    month_dir = year_dir / _month_folder_name(aook)
    project_code = (getattr(aook.project, "full_code", "") or str(aook.project or "")).strip() or "Без проекта"
    output_dir = month_dir / _safe_act_filename(project_code) / file_stem
    acts_registry_stem = _safe_filename(f"Реестр актов к АООК №{aook.number} от {aook.act_date:%d.%m.%Y}")
    protocols_registry_stem = _safe_filename(f"Реестр протоколов к АООК №{aook.number} от {aook.act_date:%d.%m.%Y}")
    return {
        "output_dir": output_dir,
        "main_docx": output_dir / f"{file_stem}.docx",
        "main_pdf": output_dir / f"{file_stem}.pdf",
        "acts_registry_docx": output_dir / f"{acts_registry_stem}.docx",
        "acts_registry_pdf": output_dir / f"{acts_registry_stem}.pdf",
        "protocols_registry_docx": output_dir / f"{protocols_registry_stem}.docx",
        "protocols_registry_pdf": output_dir / f"{protocols_registry_stem}.pdf",
    }


def get_aook_registry_pdf_path(aook: Aook, registry_type: str) -> Path:
    paths = get_aook_generated_paths(aook)
    if registry_type == "acts":
        return paths["acts_registry_pdf"]
    if registry_type == "protocols":
        return paths["protocols_registry_pdf"]
    raise AookRenderError(f"Неизвестный тип реестра АООК: {registry_type}")


def generate_and_save_aook_files(aook: Aook) -> tuple[Path, Path]:
    paths = get_aook_generated_paths(aook)
    output_dir = paths["output_dir"]
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = paths["main_docx"]
    pdf_path = paths["main_pdf"]
    acts_registry_docx_path = paths["acts_registry_docx"]
    acts_registry_pdf_path = paths["acts_registry_pdf"]
    protocols_registry_docx_path = paths["protocols_registry_docx"]
    protocols_registry_pdf_path = paths["protocols_registry_pdf"]

    render_aook_docx(aook=aook, output_path=docx_path)
    converted_pdf = convert_docx_to_pdf(docx_path, output_dir)
    if converted_pdf != pdf_path:
        shutil.move(str(converted_pdf), str(pdf_path))

    render_aook_registry_docx(aook=aook, output_path=acts_registry_docx_path, registry_type="acts")
    converted_acts_registry_pdf = convert_docx_to_pdf(acts_registry_docx_path, output_dir)
    if converted_acts_registry_pdf != acts_registry_pdf_path:
        shutil.move(str(converted_acts_registry_pdf), str(acts_registry_pdf_path))

    render_aook_registry_docx(aook=aook, output_path=protocols_registry_docx_path, registry_type="protocols")
    converted_protocols_registry_pdf = convert_docx_to_pdf(protocols_registry_docx_path, output_dir)
    if converted_protocols_registry_pdf != protocols_registry_pdf_path:
        shutil.move(str(converted_protocols_registry_pdf), str(protocols_registry_pdf_path))

    media_root = Path(settings.MEDIA_ROOT).resolve()
    aook.xlsx_file.name = ""
    try:
        aook.pdf_file.name = pdf_path.resolve().relative_to(media_root).as_posix()
    except ValueError:
        aook.pdf_file.name = ""
    aook.save(update_fields=["xlsx_file", "pdf_file", "updated_at"])
    return docx_path, pdf_path
