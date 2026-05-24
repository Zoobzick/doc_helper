# acts_app/services/registry_p3_docx_generator.py
from __future__ import annotations

import copy
import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from django.conf import settings
from django.db.models import Q
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from acts_app.models import Act, ActAttachment, AttachmentType
from acts_app.services.act_docx_context import build_projects_text
from acts_app.services.act_docx_generator import DocxRenderError, get_act_docx_paths, replace_tokens
from acts_app.services.date_format import fmt_date_g
from acts_app.services.material_resolver import resolve_material_fields
from directive_app.models import ActRole, Authorization


# -------------------------
# text utils
# -------------------------

def _safe_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _norm_spaces(s: str) -> str:
    s = (s or "").replace("\u00A0", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _uc_first(s: str) -> str:
    """Делает только первую букву заглавной (остальное не трогаем)."""
    s = _norm_spaces(s)
    if not s:
        return s
    return s[0].upper() + s[1:]


def _doc_no_with_sign(no: str) -> str:
    no = _norm_spaces(no)
    if not no or no in ("—", "-", "–"):
        return "—"
    if no.startswith("№"):
        return no
    return f"№{no}"


def _pick_registry_date(act: Act, registry: ActAttachment):
    return registry.doc_date or act.work_end_date or act.act_date


def _registry_number_prefix(registry_type: str) -> str:
    if registry_type == AttachmentType.MATERIALS_REGISTRY:
        return "П-3"
    return "Р"


def _registry_template_name(*, registry_type: str, material_rows: Iterable["_MatRow"] | None = None) -> str:
    if registry_type == AttachmentType.MATERIALS_REGISTRY:
        rows = list(material_rows or [])
        return "registry_material_concrete.docx" if _need_concrete_template(rows) else "registry_material_other.docx"
    raise DocxRenderError(f"Неизвестный тип реестра: {registry_type}")


def _registry_default_no(*, act: Act, registry_type: str) -> str:
    return f"{_registry_number_prefix(registry_type)}.{act.number}"


@dataclass(frozen=True)
class _RegistrySigner:
    organization: str
    position: str
    fio: str


def _resolve_registry_signer(act: Act) -> _RegistrySigner:
    """
    Реестры подписывает представитель лица, выполнившего работы.
    В модели акта эта сторона хранится как CONTRACTOR_REP.
    """
    party = (
        act.parties.filter(role=ActRole.CONTRACTOR_REP, is_enabled=True)
        .select_related("organization", "chosen_authorization")
        .order_by("position", "id")
        .first()
    )
    org = party.organization if (party and party.organization_id) else None
    auth = None

    if party and party.chosen_authorization_id:
        auth = (
            Authorization.objects.select_related("person")
            .filter(id=party.chosen_authorization_id)
            .first()
        )

    if auth is None and org and act.act_date:
        auth = (
            Authorization.objects.select_related("person")
            .filter(organization_id=org.id, role=ActRole.CONTRACTOR_REP, is_active=True, valid_from__lte=act.act_date)
            .filter(Q(valid_to__isnull=True) | Q(valid_to__gte=act.act_date))
            .order_by("-valid_from", "-created_at", "-id")
            .first()
        )

    person = auth.person if auth else None
    organization = (
        _strip_trailing_commas(_safe_str(getattr(org, "full_name", "")))
        or _strip_trailing_commas(_safe_str(getattr(org, "short_name", "")))
    )
    position = _strip_trailing_commas(_safe_str(getattr(auth, "position_text", "")))
    fio = (
        _safe_str(getattr(person, "short_name", ""))
        or _safe_str(getattr(person, "full_name", ""))
    )

    return _RegistrySigner(organization=organization, position=position, fio=fio)


def _safe_str(x) -> str:
    return (str(x).strip() if x is not None else "")


def _strip_trailing_commas(s: str) -> str:
    s = (s or "").strip()
    while s.endswith(","):
        s = s[:-1].rstrip()
    return s


def _registry_signer_mapping(act: Act) -> dict[str, str]:
    signer = _resolve_registry_signer(act)
    return {
        "registry_contractor_org_full": signer.organization,
        "registry_signer_position": signer.position,
        "registry_signer_fio": signer.fio,
    }


# -------------------------
# docx cell writer (preserve template formatting)
# -------------------------

def _remove_all_paragraphs_except_first(cell) -> None:
    """Удаляет лишние параграфы из ячейки, чтобы не было пустых строк/переносов."""
    if not cell.paragraphs:
        cell.add_paragraph()
        return

    # оставляем только первый параграф
    tc = cell._tc
    keep = cell.paragraphs[0]._p
    for p in list(cell.paragraphs)[1:]:
        try:
            tc.remove(p._p)
        except Exception:
            # fallback: просто очистим
            for r in p.runs:
                r.text = ""

    # убедимся что keep реально внутри
    _ = keep


def _clear_cell(cell) -> None:
    # 1) удаляем лишние параграфы
    _remove_all_paragraphs_except_first(cell)
    # 2) чистим runs в первом параграфе
    for r in cell.paragraphs[0].runs:
        r.text = ""


def _set_cell_text_from_template(cell, text: str, template_cell) -> None:
    """
    - чистим ячейку полностью (без "пустых строк"),
    - копируем стиль параграфа/рана из template_cell,
    - пишем text (если в text есть '\n' — Word покажет перенос строки).
    """
    _clear_cell(cell)
    p = cell.paragraphs[0]

    tpl_run = None
    if template_cell and template_cell.paragraphs:
        tp = template_cell.paragraphs[0]
        # стиль параграфа
        try:
            p.style = tp.style
        except Exception:
            pass
        try:
            p.paragraph_format.left_indent = tp.paragraph_format.left_indent
            p.paragraph_format.first_line_indent = tp.paragraph_format.first_line_indent
            p.paragraph_format.space_before = tp.paragraph_format.space_before
            p.paragraph_format.space_after = tp.paragraph_format.space_after
            p.paragraph_format.line_spacing = tp.paragraph_format.line_spacing
            p.paragraph_format.alignment = tp.paragraph_format.alignment
        except Exception:
            pass

        if tp.runs:
            tpl_run = tp.runs[0]

    run = p.add_run(text or "")

    if tpl_run is not None:
        run.bold = tpl_run.bold
        run.italic = tpl_run.italic
        run.underline = tpl_run.underline

        sf = tpl_run.font
        df = run.font
        df.name = sf.name
        df.size = sf.size
        if sf.color and sf.color.rgb:
            df.color.rgb = sf.color.rgb


def _compact_cell_to_single_line(cell, *, template_cell) -> None:
    """
    Убирает пустые строки/лишние параграфы из ячейки, чтобы текст не уезжал вверх.
    Оставляет одну строку (первую непустую).
    """
    raw = cell.text or ""
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    text = lines[0] if lines else ""

    _set_cell_text_from_template(cell, text, template_cell=template_cell)

    # вертикально по центру
    try:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    except Exception:
        pass


# -------------------------
# row model
# -------------------------

@dataclass(frozen=True)
class _MatRow:
    pp: int
    material_name: str
    doc_name: str
    doc_no: str
    doc_date_str: str
    pages: int
    v: int | None

    @property
    def doc_full(self) -> str:
        # {{document_name}}
        # {{document_number}} от {{document_date}}
        dn = _uc_first(self.doc_name)
        no = _doc_no_with_sign(self.doc_no)
        ds = _norm_spaces(self.doc_date_str) or "—"
        return f"{dn}\n{no} от {ds}"


# -------------------------
# docx table helpers
# -------------------------

def _insert_row_before(table, before_row_idx: int, template_row_idx: int) -> None:
    tbl = table._tbl
    before_tr = table.rows[before_row_idx]._tr
    new_tr = copy.deepcopy(table.rows[template_row_idx]._tr)
    tbl.insert(tbl.index(before_tr), new_tr)


def _append_row_clone(table, template_row_idx: int) -> None:
    new_tr = copy.deepcopy(table.rows[template_row_idx]._tr)
    table._tbl.append(new_tr)


def _merge_vertical_groups(table, *, col_idx: int, start_row: int, end_row: int, key_getter) -> None:
    """
    Объединяем подряд идущие одинаковые значения.
    ВАЖНО: перед merge очищаем нижние ячейки группы,
    иначе Word склеит текст "А + А" и получится дублирование.
    """
    if end_row <= start_row:
        return

    def key_of(r: int) -> str:
        return _norm_spaces(key_getter(r))

    group_start = start_row
    prev_key = key_of(start_row)

    for r in range(start_row + 1, end_row + 1):
        cur_key = key_of(r)
        if cur_key == prev_key:
            continue

        if r - 1 > group_start and prev_key:
            for rr in range(group_start + 1, r):
                _clear_cell(table.cell(rr, col_idx))
            table.cell(group_start, col_idx).merge(table.cell(r - 1, col_idx))

        group_start = r
        prev_key = cur_key

    if end_row > group_start and prev_key:
        for rr in range(group_start + 1, end_row + 1):
            _clear_cell(table.cell(rr, col_idx))
        table.cell(group_start, col_idx).merge(table.cell(end_row, col_idx))


# -------------------------
# collect + group rows
# -------------------------

def _collect_material_rows_for_registry(act: Act) -> list[_MatRow]:
    """
    Собираем материалы как в details (position, id),
    но ДЛЯ РЕЕСТРА группируем по material_name так, чтобы одинаковые материалы стали подряд.
    Это единственный способ объединить ячейки в Word, если одинаковые материалы разбросаны. :contentReference[oaicite:2]{index=2}
    """
    raw: list[_MatRow] = []

    qs = act.materials.select_related("passport", "passport__material").order_by("position", "id")
    for _, mi in enumerate(qs, start=1):
        data = resolve_material_fields(mi)

        material_name = _uc_first(data.get("material_name") or "Материал")
        doc_name = _uc_first(data.get("document_name") or "Документ")
        doc_no = _norm_spaces(data.get("document_no") or "—")
        doc_date_str = _norm_spaces(data.get("document_date_str") or "—")

        raw.append(
            _MatRow(
                pp=0,  # проставим после группировки
                material_name=material_name,
                doc_name=doc_name,
                doc_no=doc_no,
                doc_date_str=doc_date_str,
                pages=int(mi.sheets_count or 0),
                v=(int(mi.concrete_volume_m3) if getattr(mi, "concrete_volume_m3", None) not in (None, "") else None),
            )
        )

    # ✅ Группируем по material_name, сохраняя порядок первого появления материала
    buckets: dict[str, list[_MatRow]] = {}
    order: list[str] = []
    for r in raw:
        key = _norm_spaces(r.material_name).lower()
        if key not in buckets:
            buckets[key] = []
            order.append(key)
        buckets[key].append(r)

    grouped: list[_MatRow] = []
    for key in order:
        grouped.extend(buckets[key])

    # перенумерация pp по итоговому порядку в реестре
    out: list[_MatRow] = []
    for i, r in enumerate(grouped, start=1):
        out.append(_MatRow(pp=i, **{k: getattr(r, k) for k in r.__dataclass_fields__ if k != "pp"}))

    return out


def _need_concrete_template(rows: Iterable[_MatRow]) -> bool:
    return any(r.v is not None for r in rows)


# -------------------------
# public API
# -------------------------

def get_registry_p3_docx_paths(*, act: Act, registry: ActAttachment) -> list[Path]:
    act_paths = get_act_docx_paths(act)
    if not act_paths:
        raise DocxRenderError("Не удалось определить путь акта (get_act_docx_paths вернул пусто).")

    d = _pick_registry_date(act, registry)
    date_str = d.strftime("%d.%m.%Y") if d else "—"

    reg_no = (registry.doc_no or _registry_default_no(act=act, registry_type=AttachmentType.MATERIALS_REGISTRY)).strip()
    file_name = _safe_filename(f"Реестр №{reg_no} от {date_str}") + ".docx"

    return [p.parent / file_name for p in act_paths]


def generate_registry_p3_docx_bytes(*, act: Act, registry: ActAttachment) -> bytes:
    rows = _collect_material_rows_for_registry(act)
    if not rows:
        raise DocxRenderError("Нельзя собрать реестр П-3: в акте нет материалов.")

    tpl_name = _registry_template_name(registry_type=AttachmentType.MATERIALS_REGISTRY, material_rows=rows)
    template_path = Path(settings.DOCX_TEMPLATES_DIR) / tpl_name
    if not template_path.exists():
        raise DocxRenderError(f"Не найден DOCX-шаблон реестра П-3: {template_path}")

    doc = Document(str(template_path))

    proj = build_projects_text(act)
    mapping = {
        "project_line": (proj.get("project_line") or "").strip(),
        "project_stage": (proj.get("project_stage") or "").strip(),
        "project_address": (proj.get("project_address") or "").strip(),
        "projects_full_code": ", ".join(
            [(getattr(p, "full_code", "") or "").strip() or str(p).strip() for p in act.projects.all().order_by("id")]
        ).strip() or "—",
        "registry_name": (registry.doc_no or _registry_default_no(act=act, registry_type=AttachmentType.MATERIALS_REGISTRY)).strip(),
        "registry_date": fmt_date_g(_pick_registry_date(act, registry)),
        "act_name": (act.number or "").strip(),
        "act_date": fmt_date_g(act.act_date),
    }
    mapping.update(_registry_signer_mapping(act))
    replace_tokens(doc, mapping)

    if not doc.tables:
        raise DocxRenderError("В шаблоне реестра П-3 не найдена таблица.")
    table = doc.tables[0]

    if len(table.rows) < 2:
        raise DocxRenderError("Таблица реестра П-3 должна иметь минимум 2 строки (header + template).")

    template_idx = 1
    is_concrete = _need_concrete_template(rows)

    total_idx = 2 if is_concrete else None
    if is_concrete and len(table.rows) < 3:
        raise DocxRenderError("Concrete-шаблон П-3 должен иметь строку итога (3-я строка).")

    # расширяем таблицу, сохраняя форматирование строк
    if is_concrete:
        for _ in range(len(rows) - 1):
            _insert_row_before(table, before_row_idx=total_idx, template_row_idx=template_idx)
        total_idx = 1 + len(rows)
    else:
        for _ in range(len(rows) - 1):
            _append_row_clone(table, template_row_idx=template_idx)

    tpl_cells = table.rows[template_idx].cells

    # заполняем строки
    for i, r in enumerate(rows):
        row = table.rows[template_idx + i]

        _set_cell_text_from_template(row.cells[0], str(r.pp), template_cell=tpl_cells[0])
        _set_cell_text_from_template(row.cells[1], _uc_first(r.material_name), template_cell=tpl_cells[1])
        _set_cell_text_from_template(row.cells[2], r.doc_full, template_cell=tpl_cells[2])

        if is_concrete:
            v_txt = str(r.v) if r.v is not None else "-"
            _set_cell_text_from_template(row.cells[3], v_txt, template_cell=tpl_cells[3])
            _set_cell_text_from_template(row.cells[4], str(r.pages or 0), template_cell=tpl_cells[4])
        else:
            _set_cell_text_from_template(row.cells[3], str(r.pages or 0), template_cell=tpl_cells[3])

    # sum_concrete_v
    if is_concrete and total_idx is not None:
        sum_v = sum((r.v or 0) for r in rows)
        total_row = table.rows[total_idx]
        _set_cell_text_from_template(total_row.cells[3], str(sum_v), template_cell=total_row.cells[3])

    # объединения (теперь одинаковые material_name уже стоят подряд благодаря группировке)
    start = template_idx
    end = template_idx + len(rows) - 1

    _merge_vertical_groups(
        table,
        col_idx=1,
        start_row=start,
        end_row=end,
        key_getter=lambda ri: table.cell(ri, 1).text,
    )

    if not is_concrete:
        _merge_vertical_groups(
            table,
            col_idx=2,
            start_row=start,
            end_row=end,
            key_getter=lambda ri: table.cell(ri, 2).text,
        )
    for ri in range(start, end + 1):
        _compact_cell_to_single_line(table.cell(ri, 1), template_cell=tpl_cells[1])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_and_save_registry_p3_docx(*, act: Act, registry: ActAttachment) -> list[Path]:
    if registry.type != AttachmentType.MATERIALS_REGISTRY:
        raise DocxRenderError("generate_and_save_registry_p3_docx: registry.type должен быть MATERIALS_REGISTRY.")

    content = generate_registry_p3_docx_bytes(act=act, registry=registry)
    paths = get_registry_p3_docx_paths(act=act, registry=registry)

    saved: list[Path] = []
    for p in paths:
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_bytes(content)
        saved.append(p)

    return saved
